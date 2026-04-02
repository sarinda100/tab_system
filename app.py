import sys
import subprocess
try:
    import openpyxl
except ImportError:
    print("Installing openpyxl automatically...")
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'openpyxl'])
from docx import Document
from docx.shared import Inches, Pt
import io
import json
import csv
import os
import re
import pymysql
import pymysql.cursors
import openpyxl  # ADDED: Excel Support
from openpyxl.styles import Font # ADDED: For Template styling
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, KeepTogether
from io import BytesIO
import pytz
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, Response, session, send_file

from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
import os
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY')

@app.before_request
def check_maintenance_mode():
    if os.path.exists('update.txt'):
        return """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>System Maintenance</title>
            <style>
                body { background: #080f1c; color: #ddeeff; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; display: flex; align-items: center; justify-content: center; height: 100vh; margin: 0; text-align: center; }
                .box { background: #0f2040; padding: 40px; border-radius: 16px; border: 1px solid #00d4ff; box-shadow: 0 10px 30px rgba(0,212,255,0.1); max-width: 400px;}
                h1 { color: #00d4ff; margin-bottom: 10px; font-size: 24px;}
                p { color: #4a6a8a; font-size: 14px; line-height: 1.6; }
                .loader { border: 4px solid #1a3050; border-top: 4px solid #00d4ff; border-radius: 50%; width: 40px; height: 40px; animation: spin 1s linear infinite; margin: 0 auto 20px auto; }
                @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
            </style>
        </head>
        <body>
            <div class="box">
                <div class="loader"></div>
                <h1>🛠️ System Updating</h1>
                <p>TABCORE is currently undergoing scheduled maintenance and upgrades. <br><br>Please check back in a few minutes. Your data is safe!</p>
            </div>
        </body>
        </html>
        """, 503

app.config['REMEMBER_COOKIE_DURATION'] = timedelta(days=30) 

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

def get_db_connection():
    return pymysql.connect(
        host=os.getenv('DB_HOST'),
        user=os.getenv('DB_USER'),
        password=os.getenv('DB_PASSWORD'),
        database=os.getenv('DB_NAME'),
        cursorclass=pymysql.cursors.DictCursor
    )

class User(UserMixin):
    def __init__(self, id, username, name, role):
        self.id = id
        self.username = username
        self.name = name
        self.role = role

@login_manager.user_loader
def load_user(user_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
    user = cursor.fetchone()
    conn.close()
    if user: return User(user['id'], user['username'], user['name'], user['role'])
    return None

# --- STRICT BACKEND VALIDATIONS & EXCEL FIXES ---

def clean_imei(imei_str):
    if not imei_str: return ""
    imei_str = str(imei_str).strip().upper()
    if 'E+' in imei_str:
        try: return "{:.0f}".format(float(imei_str))
        except: return imei_str
    return re.sub(r'[^0-9]', '', imei_str)

def is_valid_imei(imei_str):
    if not imei_str: return False
    return bool(re.fullmatch(r'^\d{15}$', str(imei_str)))

def format_and_validate_sn(sn_str, brand=""):
    if not sn_str: return None
    clean_sn = str(sn_str).replace(" ", "").upper()
    if 'E+' in clean_sn:
        try: clean_sn = "{:.0f}".format(float(clean_sn))
        except: pass
    clean_sn = re.sub(r'[^A-Z0-9]', '', clean_sn)
    
    brand_lower = str(brand).lower()
    
    if 'samsung' in brand_lower:
        if len(clean_sn) == 11:
            return clean_sn
        return None
    elif 'lenovo' in brand_lower:
        if len(clean_sn) == 8:
            return clean_sn
        return None
    else:
        if 8 <= len(clean_sn) <= 15:
            return clean_sn
        return None

def is_valid_asset_no(asset_str):
    if not asset_str: return True 
    return bool(re.fullmatch(r'^\d{5}$', str(asset_str)))

# --- HELPERS ---

def log_device_history(tablet_id, action, performed_by, status_changed_to=None, notes=None):
    conn = get_db_connection(); cursor = conn.cursor()
    cursor.execute("INSERT INTO device_history (tablet_id, action, performed_by, status_changed_to, notes) VALUES (%s, %s, %s, %s, %s)", (tablet_id, action, performed_by, status_changed_to, notes))
    conn.commit(); conn.close()

def log_system_audit(action, performed_by, details=None):
    conn = get_db_connection(); cursor = conn.cursor()
    cursor.execute("INSERT INTO system_audit (action, performed_by, details) VALUES (%s, %s, %s)", (action, performed_by, details))
    conn.commit(); conn.close()

def safe_excel_text(val):
    if not val or val == '-': return '-'
    return f'="{val}"'

# --- DATABASE INIT ---

def create_tables():
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("ALTER TABLE tablets ADD COLUMN pouch_status VARCHAR(50) DEFAULT 'Missing'")
        cursor.execute("ALTER TABLE tablets ADD COLUMN pen_status VARCHAR(50) DEFAULT 'Missing'")
        conn.commit()
    except: pass 

    cursor.execute("CREATE TABLE IF NOT EXISTS device_history (id INT AUTO_INCREMENT PRIMARY KEY, tablet_id INT NOT NULL, action VARCHAR(50) NOT NULL, performed_by VARCHAR(100) NOT NULL, status_changed_to VARCHAR(50), notes TEXT, timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP, FOREIGN KEY (tablet_id) REFERENCES tablets(id) ON DELETE CASCADE)")
    cursor.execute("CREATE TABLE IF NOT EXISTS system_audit (id INT AUTO_INCREMENT PRIMARY KEY, action VARCHAR(100) NOT NULL, performed_by VARCHAR(100) NOT NULL, details TEXT, timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
    cursor.execute("CREATE TABLE IF NOT EXISTS settings (id INT PRIMARY KEY, batch_target INT DEFAULT 540)")
    cursor.execute("CREATE TABLE IF NOT EXISTS district_targets (district_name VARCHAR(100) PRIMARY KEY, target_count INT DEFAULT 0)")
    
    valid_districts = ['Colombo', 'Gampaha', 'Batticaloa', 'Trincomalee', 'Anuradhapura', 'Polonnaruwa', 'Badulla', 'Ratnapura', 'Kegalle']
    format_strings = ','.join(['%s'] * len(valid_districts))
    cursor.execute(f"DELETE FROM district_targets WHERE district_name NOT IN ({format_strings})", valid_districts)

    cursor.execute("SELECT COUNT(*) as c FROM district_targets")
    if cursor.fetchone()['c'] == 0:
        default_districts = [('Colombo', 108), ('Gampaha', 77), ('Batticaloa', 54), ('Trincomalee', 40), ('Anuradhapura', 7), ('Polonnaruwa', 12), ('Badulla', 81), ('Ratnapura', 106), ('Kegalle', 58)]
        cursor.executemany("INSERT INTO district_targets (district_name, target_count) VALUES (%s, %s)", default_districts)

    cursor.execute("INSERT IGNORE INTO settings (id, batch_target) VALUES (1, 540)")
    conn.commit(); conn.close()

def create_default_admin():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE username='admin'")
    if not cursor.fetchone():
        hashed_pw = generate_password_hash('admin123')
        cursor.execute("INSERT INTO users (username, password, name, role) VALUES (%s, %s, %s, %s)", ('admin', hashed_pw, 'Sarinda', 'Admin'))
        conn.commit()
    conn.close()

# --- AUTH ROUTES ---

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated: return redirect(url_for('dashboard'))
    error = None
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        remember_me = request.form.get('remember') == 'yes' 

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        conn.close()
        
        if user and check_password_hash(user['password'], password):
            user_obj = User(user['id'], user['username'], user['name'], user['role'])
            login_user(user_obj, remember=remember_me)
            log_system_audit("User Login", user['name'], f"User '{username}' logged in.")
            return redirect(url_for('dashboard'))
        error = "Invalid username or password"
    return render_template('login.html', error=error)

@app.route('/logout')
@login_required
def logout():
    log_system_audit("User Logout", current_user.name, f"User '{current_user.username}' logged out.")
    logout_user()
    return redirect(url_for('login'))

@app.route('/')
def index():
    return redirect(url_for('dashboard'))

# --- DASHBOARD ---

@app.route('/dashboard')
@login_required
def dashboard():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT COUNT(*) as total FROM tablets WHERE is_deleted = 0")
    total = cursor.fetchone()['total']
    
    cursor.execute("SELECT COUNT(*) as pending FROM tablets WHERE status IN ('Pending', 'Locked') AND is_deleted = 0")
    pending = cursor.fetchone()['pending']
    
    cursor.execute("SELECT COUNT(*) as inspected FROM tablets WHERE status IN ('Passed', 'Minor Issues', 'Defective', 'Dead Device (DOA)') AND is_deleted = 0")
    inspected = cursor.fetchone()['inspected']
    
    cursor.execute("SELECT COUNT(*) as defective FROM tablets WHERE status = 'Defective' AND is_deleted = 0")
    defective = cursor.fetchone()['defective']
    
    cursor.execute("SELECT COUNT(*) as doa_count FROM tablets WHERE status = 'Dead Device (DOA)' AND is_deleted = 0")
    doa_count = cursor.fetchone()['doa_count']
    
    cursor.execute("SELECT * FROM tablets WHERE is_deleted = 0 ORDER BY id DESC LIMIT 5")
    recent = cursor.fetchall()
    
    cursor.execute("SELECT batch_target FROM settings WHERE id = 1")
    res = cursor.fetchone(); batch_target = res['batch_target'] if res else 540
    pipeline_percentage = round((inspected / batch_target) * 100, 1) if batch_target > 0 else 0

    cursor.execute("SELECT COUNT(*) as c FROM tablets WHERE DATE(registered_at) = CURDATE()"); today_total = cursor.fetchone()['c']
    cursor.execute("SELECT COUNT(DISTINCT tablet_id) as c FROM device_history WHERE action = 'Inspected' AND DATE(timestamp) = CURDATE()"); today_inspected = cursor.fetchone()['c']

    cursor.execute("SELECT COUNT(*) as c FROM tablets WHERE LOWER(brand) LIKE '%samsung%'"); samsung_count = cursor.fetchone()['c']
    cursor.execute("SELECT COUNT(*) as c FROM tablets WHERE LOWER(brand) LIKE '%lenovo%'"); lenovo_count = cursor.fetchone()['c']
    cursor.execute("SELECT COUNT(*) as c FROM tablets WHERE LOWER(charger_status) LIKE '%missing%' OR LOWER(cable_status) LIKE '%missing%' OR LOWER(simpin_status) LIKE '%missing%'"); missing_acc = cursor.fetchone()['c']
    cursor.execute("SELECT COUNT(DISTINCT inspected_by) as c FROM tablets WHERE inspected_by IS NOT NULL AND inspected_by != ''"); active_techs = cursor.fetchone()['c']

    cursor.execute("SELECT u.name as inspected_by, COUNT(t.id) as count FROM users u LEFT JOIN tablets t ON u.name = t.inspected_by AND t.status IN ('Passed', 'Minor Issues', 'Defective', 'Dead Device (DOA)') GROUP BY u.name")
    user_data = cursor.fetchall(); user_labels = [row['inspected_by'] for row in user_data]; user_counts = [row['count'] for row in user_data]

    cursor.execute("SELECT DATE(registered_at) as date_val, COUNT(id) as count FROM tablets WHERE status IN ('Passed', 'Minor Issues', 'Defective', 'Dead Device (DOA)') GROUP BY DATE(registered_at) ORDER BY date_val DESC LIMIT 7")
    daily_data = cursor.fetchall(); daily_labels = [str(row['date_val']) for row in daily_data]; daily_counts = [row['count'] for row in daily_data]
    daily_labels.reverse(); daily_counts.reverse()

    cursor.execute("""
        SELECT dt.district_name, dt.target_count, COUNT(t.id) as inspected_count
        FROM district_targets dt LEFT JOIN tablets t ON dt.district_name = t.district AND t.status IN ('Passed', 'Minor Issues', 'Defective', 'Dead Device (DOA)')
        WHERE dt.target_count > 0 GROUP BY dt.district_name, dt.target_count ORDER BY (dt.target_count - COUNT(t.id)) DESC
    """)
    district_progress_data = cursor.fetchall()

    district_stats = []
    for dp in district_progress_data:
        target = dp['target_count']; inspected_count = dp['inspected_count']; remaining = max(0, target - inspected_count); percentage = min(100, round((inspected_count / target) * 100)) if target > 0 else 0
        district_stats.append({'name': dp['district_name'], 'target': target, 'inspected': inspected_count, 'remaining': remaining, 'percentage': percentage})

    conn.close()
    
    return render_template('dashboard.html', total_tablets=total, pending=pending, inspected=inspected, defective=defective, doa_count=doa_count, today_total=today_total, today_inspected=today_inspected, samsung_count=samsung_count, lenovo_count=lenovo_count, missing_acc=missing_acc, active_techs=active_techs, batch_target=batch_target, pipeline_percentage=pipeline_percentage, recent_tablets=recent, user_labels=json.dumps(user_labels), user_counts=json.dumps(user_counts), daily_labels=json.dumps(daily_labels), daily_counts=json.dumps(daily_counts), district_stats=district_stats)

@app.route('/register', methods=['GET', 'POST'])
@login_required
def register():
    error = None
    if request.method == 'POST':
        if 'file' in request.files and request.files['file'].filename != '':
            file = request.files['file']
            
            if not (file.filename.endswith('.csv') or file.filename.endswith('.xlsx')):
                return render_template('register.html', error="Please upload a valid CSV (.csv) or Excel (.xlsx) file.")
            
            try:
                rows_data = []
                
                if file.filename.endswith('.csv'):
                    content = file.stream.read()
                    try: decoded_content = content.decode('utf-8-sig')
                    except: decoded_content = content.decode('cp1252', errors='ignore')
                    csv_input = csv.DictReader(io.StringIO(decoded_content, newline=None))
                    for raw_row in csv_input:
                        row = {str(k).strip(): str(v).strip() for k, v in raw_row.items() if k}
                        rows_data.append(row)
                
                elif file.filename.endswith('.xlsx'):
                    wb = openpyxl.load_workbook(file, data_only=True)
                    ws = wb.active
                    sheet_rows = list(ws.rows)
                    if len(sheet_rows) > 0:
                        headers = [str(cell.value).strip() if cell.value is not None else '' for cell in sheet_rows[0]]
                        for sheet_row in sheet_rows[1:]:
                            row_dict = {}
                            for i, cell in enumerate(sheet_row):
                                if i < len(headers) and headers[i]:
                                    val = cell.value
                                    if val is None:
                                        val_str = ""
                                    elif isinstance(val, float) and val.is_integer():
                                        val_str = str(int(val))
                                    else:
                                        val_str = str(val)
                                    row_dict[headers[i]] = val_str.strip()
                            rows_data.append(row_dict)

                conn = get_db_connection()
                cursor = conn.cursor()
                success_count = 0
                duplicate_count = 0
                invalid_imei_count = 0
                invalid_sn_count = 0
                invalid_asset_count = 0
                
                for row in rows_data:
                    raw_sn = row.get('Serial Number', row.get('S/N', ''))
                    raw_imei = row.get('IMEI Number', row.get('IMEI', ''))
                    raw_asset = str(row.get('Asset No', row.get('Asset Number', ''))).strip() 
                    if raw_asset == 'None': raw_asset = ''
                    if not raw_sn or not raw_imei: continue 
                    
                    imei = clean_imei(raw_imei)
                    if not is_valid_imei(imei):
                        invalid_imei_count += 1
                        continue

                    sn = format_and_validate_sn(raw_sn, row.get('Brand', ''))
                    if not sn:
                        invalid_sn_count += 1
                        continue
                        
                    asset_no_clean = re.sub(r'[^0-9]', '', raw_asset)
                    if raw_asset and not is_valid_asset_no(asset_no_clean):
                        invalid_asset_count += 1
                        continue

                    cursor.execute("SELECT id FROM tablets WHERE serial_number = %s OR imei_number = %s", (sn, imei))
                    if cursor.fetchone():
                        duplicate_count += 1
                        continue 
                        
                    sql = "INSERT INTO tablets (district, brand, model, serial_number, asset_no, imei_number, charger_status, cable_status, simpin_status, pouch_status, pen_status, doc_status, registered_by, status) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'Bulk','Bulk',%s,%s,'Pending')" 
                    cursor.execute(sql, (row.get('District', ''), row.get('Brand', ''), row.get('Model', ''), sn, asset_no_clean, imei, row.get('Charger Status') or 'Good', row.get('Cable Status') or 'Good', row.get('Simpin Status') or 'Good', row.get('Doc Status') or 'Good', current_user.name))
                    tid = cursor.lastrowid
                    conn.commit()
                    log_device_history(tid, "Registered (Bulk)", current_user.name, "Pending", "Bulk registered via File Upload")
                    success_count += 1
                    
                conn.close()
                msg = f"Bulk Upload Complete: {success_count} added. {duplicate_count} duplicates skipped."
                if invalid_imei_count > 0: msg += f" {invalid_imei_count} skipped due to INVALID IMEI."
                if invalid_sn_count > 0: msg += f" {invalid_sn_count} skipped due to INVALID S/N."
                if invalid_asset_count > 0: msg += f" {invalid_asset_count} skipped due to INVALID ASSET NO."
                flash(msg, "success" if (invalid_imei_count == 0 and invalid_sn_count == 0 and invalid_asset_count == 0) else "error")
                return redirect(url_for('dashboard'))
            except Exception as e:
                error = f"Error processing Upload file: {str(e)}"

        else:
            d = request.form
            raw_sn = d.get('serial_number', '').strip()
            raw_imei = d.get('imei_number', '').strip()
            asset_no = d.get('asset_no', '').strip() 
            
            if raw_sn and raw_imei:
                imei = clean_imei(raw_imei)
                if not is_valid_imei(imei):
                    return render_template('register.html', error="Validation Error: IMEI Number must be EXACTLY 15 digits! No letters or spaces allowed.")

                try:
                    brand = d.get('brand', '')
                    sn = format_and_validate_sn(raw_sn, brand)
                    
                    if not sn:
                        if 'samsung' in brand.lower():
                            error = "Validation Error: Samsung Serial Number must be EXACTLY 11 characters!"
                        elif 'lenovo' in brand.lower():
                            error = "Validation Error: Lenovo Serial Number must be EXACTLY 8 characters!"
                        else:
                            error = "Validation Error: Serial Number must be between 8-15 characters."
                    
                    elif asset_no and (len(asset_no) != 5 or not asset_no.isdigit()):
                        error = f"Error: Asset Number '{asset_no}' is invalid. It must be exactly 5 digits (e.g., 12345)!"
                    
                    else:
                        conn = get_db_connection()
                        cursor = conn.cursor()
                        
                        cursor.execute("""
                            SELECT serial_number, imei_number, asset_no 
                            FROM tablets 
                            WHERE serial_number = %s OR imei_number = %s OR (asset_no = %s AND asset_no != '')
                        """, (sn, imei, asset_no))
                        
                        existing_device = cursor.fetchone()

                        if existing_device:
                            if existing_device['serial_number'] == sn:
                                error = f"Error: Serial Number '{sn}' is already registered in the system!"
                            elif existing_device['imei_number'] == imei:
                                error = f"Error: IMEI Number '{imei}' is already registered!"
                            elif asset_no and existing_device['asset_no'] == asset_no:
                                error = f"Error: Asset Number '{asset_no}' is already assigned to another device!"
                        else:
                            sql = "INSERT INTO tablets (district, brand, model, serial_number, asset_no, imei_number, charger_status, cable_status, simpin_status, pouch_status, pen_status, doc_status, registered_by, status) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'Bulk','Bulk',%s,%s,'Pending')" 
                            cursor.execute(sql, (d['district'], brand, d.get('model',''), sn, asset_no, imei, d['charger_status'], d['cable_status'], d['simpin_status'], d['doc_status'], current_user.name))
                            tid = cursor.lastrowid
                            conn.commit()
                            log_device_history(tid, "Registered", current_user.name, "Pending", "Manual Registration")
                            flash(f"Device {sn} registered successfully!", "success")
                            return redirect(url_for('dashboard'))
                except Exception as e: 
                    error = f"Database Error: {str(e)}"
                finally: 
                    if 'conn' in locals() and conn.open:
                        conn.close()

    return render_template('register.html', error=error)

# --- INSPECTION QUEUE ---

@app.route('/inspection')
@login_required
def inspection_queue():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE tablets SET status = 'Pending' WHERE status = 'Locked' AND registered_at <= DATE_SUB(NOW(), INTERVAL 2 DAY)")
    conn.commit()
    cursor.execute("SELECT * FROM tablets WHERE status IN ('Pending', 'Locked') AND is_deleted = 0 ORDER BY id ASC")
    tablets = cursor.fetchall()
    conn.close()
    return render_template('queue.html', tablets=tablets)

@app.route('/inspect/<int:id>', methods=['GET', 'POST'])
@login_required
def inspect_device(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT status, inspected_by FROM tablets WHERE id = %s", (id,))
    tablet_check = cursor.fetchone()
    if not tablet_check:
        conn.close()
        return redirect(url_for('inspection_queue'))
    if tablet_check['status'] == 'Locked' and tablet_check['inspected_by'] != current_user.name:
        conn.close()
        flash(f"Access Denied! This device is currently being inspected by {tablet_check['inspected_by']}.", "error")
        return redirect(url_for('inspection_queue'))
    if request.method == 'POST':
        verdict = request.form.get('verdict')
        battery_drain_time = request.form.get('battery_drain_time', '-')
        checklist = {k: v for k, v in request.form.items() if k not in ['verdict', 'battery_drain_time']}
        cursor.execute("UPDATE tablets SET status = %s, inspected_by = %s, inspection_data = %s, battery_drain_time = %s WHERE id = %s", (verdict, current_user.name, json.dumps(checklist), battery_drain_time, id))
        conn.commit()
        log_device_history(id, "Inspected", current_user.name, verdict, f"Verdict: {verdict}")
        conn.close()
        return redirect(url_for('inspection_queue'))
    if tablet_check['status'] == 'Pending':
        cursor.execute("UPDATE tablets SET status = 'Locked', inspected_by = %s WHERE id = %s", (current_user.name, id))
        conn.commit()
        log_device_history(id, "Locked for Inspection", current_user.name, "Locked", "Technician opened for inspection")
    cursor.execute("SELECT * FROM tablets WHERE id = %s", (id,))
    tablet = cursor.fetchone()
    conn.close()
    return render_template('inspect.html', tablet=tablet)

@app.route('/verified')
@login_required
def verified_records():
    page = request.args.get('page', 1, type=int)
    per_page = 50
    offset = (page - 1) * per_page

    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT COUNT(*) as total FROM tablets WHERE status IN ('Passed', 'Minor Issues') AND is_deleted = 0")
    total_records = cursor.fetchone()['total']
    total_pages = (total_records + per_page - 1) // per_page
    if total_pages == 0: total_pages = 1

    cursor.execute("SELECT * FROM tablets WHERE status IN ('Passed', 'Minor Issues') AND is_deleted = 0 ORDER BY id DESC LIMIT %s OFFSET %s", (per_page, offset))
    tablets = cursor.fetchall()
    conn.close()
    
    return render_template('records.html', tablets=tablets, title="Verified Records", color="var(--green)", page=page, total_pages=total_pages)

@app.route('/defects')
@login_required
def defect_records():
    page = request.args.get('page', 1, type=int)
    per_page = 50
    offset = (page - 1) * per_page

    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT COUNT(*) as total FROM tablets WHERE status = 'Defective' AND is_deleted = 0")
    total_records = cursor.fetchone()['total']
    total_pages = (total_records + per_page - 1) // per_page
    if total_pages == 0: total_pages = 1

    cursor.execute("SELECT * FROM tablets WHERE status = 'Defective' AND is_deleted = 0 ORDER BY id DESC LIMIT %s OFFSET %s", (per_page, offset))
    tablets = cursor.fetchall()
    conn.close()
    
    return render_template('records.html', tablets=tablets, title="Issues / Defects", color="var(--red)", page=page, total_pages=total_pages)

@app.route('/accessories')
@login_required
def accessories_log():
    page = request.args.get('page', 1, type=int)
    per_page = 50
    offset = (page - 1) * per_page

    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT COUNT(*) as chargers FROM tablets WHERE LOWER(charger_status) != 'missing' AND is_deleted = 0")
    chargers = cursor.fetchone()['chargers']
    cursor.execute("SELECT COUNT(*) as cables FROM tablets WHERE LOWER(cable_status) != 'missing' AND is_deleted = 0")
    cables = cursor.fetchone()['cables']
    cursor.execute("SELECT COUNT(*) as simpins FROM tablets WHERE LOWER(simpin_status) != 'missing' AND is_deleted = 0")
    simpins = cursor.fetchone()['simpins']
    cursor.execute("SELECT COUNT(*) as pouches FROM tablets WHERE LOWER(pouch_status) != 'missing' AND is_deleted = 0")
    pouches = cursor.fetchone()['pouches']
    cursor.execute("SELECT COUNT(*) as pens FROM tablets WHERE LOWER(pen_status) != 'missing' AND is_deleted = 0")
    pens = cursor.fetchone()['pens']
    
    cursor.execute("SELECT COUNT(*) as total FROM tablets WHERE is_deleted = 0")
    total_records = cursor.fetchone()['total']
    total_pages = (total_records + per_page - 1) // per_page
    if total_pages == 0: total_pages = 1
    
    cursor.execute("SELECT id, serial_number, asset_no, brand, district, charger_status, cable_status, simpin_status, pouch_status, pen_status, registered_by FROM tablets WHERE is_deleted = 0 ORDER BY id DESC LIMIT %s OFFSET %s", (per_page, offset))
    tablets = cursor.fetchall()    
    conn.close()
    
    return render_template('accessories.html', tablets=tablets, chargers=chargers, cables=cables, simpins=simpins, pouches=pouches, pens=pens, page=page, total_pages=total_pages)

# --- EXPORT ROUTES ---

@app.route('/export_accessories')
@login_required
def export_accessories():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, district, brand, serial_number, asset_no, charger_status, cable_status, simpin_status, pouch_status, pen_status, doc_status, registered_by FROM tablets ORDER BY id ASC")
    tablets = cursor.fetchall()
    conn.close()
    
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Tablet ID', 'District', 'Brand', 'Serial Number', 'Asset No', 'Charger', 'Data Cable', 'SIM Pin', 'Rugged Pouch', 'Stylus Pen', 'Delivery Note', 'Registered By'])
    
    for t in tablets:
        safe_sn = safe_excel_text(t.get('serial_number'))
        writer.writerow([
            f"T-{t['id']}", t.get('district','-'), t.get('brand','-'), safe_sn, t.get('asset_no', '-'),
            t.get('charger_status','-'), t.get('cable_status','-'), t.get('simpin_status','-'),
            t.get('pouch_status','-'), t.get('pen_status','-'), t.get('doc_status','-'), t.get('registered_by','-')
        ])
        
    log_system_audit("Accessories Export", current_user.name, "Exported accessories detailed report to CSV.")
    return Response(output.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment;filename=Tabcore_Accessories_Report.csv"})

@app.route('/export')
@login_required
def export_data():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))

    try:
        # Bulletproof user capture
        safe_user_name = "System Admin"
        try:
            if current_user:
                if hasattr(current_user, 'name') and current_user.name:
                    safe_user_name = current_user.name
                elif hasattr(current_user, 'username') and current_user.username:
                    safe_user_name = current_user.username
        except: pass

        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.chart import PieChart, Reference
        from openpyxl.chart.series import DataPoint
        from datetime import datetime
        import pytz, io, json

        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Fetching entire inventory
        cursor.execute("SELECT * FROM tablets WHERE is_deleted = 0 ORDER BY district ASC, id ASC")
        tablets = cursor.fetchall()
        conn.close()

        if not tablets:
            return "No records found to export."

        tz = pytz.timezone('Asia/Colombo')
        display_time = datetime.now(tz).strftime('%Y-%m-%d %I:%M %p')

        # --- STYLES ---
        def thin_border():
            s = Side(style='thin', color='CCCCCC')
            return Border(left=s, right=s, top=s, bottom=s)

        navy_fill   = PatternFill("solid", fgColor="1A2A6C")
        pass_fill   = PatternFill("solid", fgColor="D4EDDA")
        minor_fill  = PatternFill("solid", fgColor="FFF3CD")
        fail_fill   = PatternFill("solid", fgColor="F8D7DA")
        doa_fill    = PatternFill("solid", fgColor="EDE9FE")
        alt_fill    = PatternFill("solid", fgColor="F8FAFF")

        STATUS_FILL = {
            'Passed': pass_fill,
            'Minor Issues': minor_fill,
            'Defective': fail_fill,
            'Dead Device (DOA)': doa_fill,
        }

        # Pre-compute Stats
        stats = {'total': len(tablets), 'passed': 0, 'minor': 0, 'defective': 0, 'doa': 0, 'pending': 0}
        district_stats = {}
        shortages_list = [] 

        item_keys = ['display','touch','battery','cameras','wifi','bt','gps','speaker','mic','charging','p_btn','sim']
        ITEM_LABELS = ['Display','Touch','Battery','Camera','Wi-Fi','Bluetooth','GPS','Speaker','Mic','Charging Port','Power Btn','SIM Slot']

        for t in tablets:
            st = str(t.get('status',''))
            d = str(t.get('district') or 'Unknown').strip().title()

            if st == 'Passed': stats['passed'] += 1
            elif st == 'Minor Issues': stats['minor'] += 1
            elif st == 'Defective': stats['defective'] += 1
            elif st == 'Dead Device (DOA)': stats['doa'] += 1
            else: stats['pending'] += 1

            if d not in district_stats:
                district_stats[d] = {'total':0, 'passed':0, 'minor':0, 'defective':0, 'doa':0}
            
            district_stats[d]['total'] += 1
            if st == 'Passed': district_stats[d]['passed'] += 1
            elif st == 'Minor Issues': district_stats[d]['minor'] += 1
            elif st == 'Defective': district_stats[d]['defective'] += 1
            elif st == 'Dead Device (DOA)': district_stats[d]['doa'] += 1

            # Accessory shortages (🔴 Updated to include both Missing & Damaged)
            cs = str(t.get('charger_status','')).lower()
            cb = str(t.get('cable_status','')).lower()
            sp = str(t.get('simpin_status','')).lower()
            
            is_cs_bad = 'missing' in cs or 'damage' in cs
            is_cb_bad = 'missing' in cb or 'damage' in cb
            is_sp_bad = 'missing' in sp or 'damage' in sp
            
            if is_cs_bad or is_cb_bad or is_sp_bad:
                shortages_list.append({
                    'id': t['id'], 'district': d, 'serial': t.get('serial_number','-'),
                    'charger': t.get('charger_status','-'), 'cable': t.get('cable_status','-'),
                    'sim_pin': t.get('simpin_status','-'), 'status': st
                })

        wb = Workbook()

        # =========================================================
        # SHEET 1: MASTER DASHBOARD
        # =========================================================
        ws_s = wb.active
        ws_s.title = "📊 Master Dashboard"
        ws_s.sheet_properties.tabColor = "1A2A6C"
        ws_s.sheet_view.showGridLines = False

        for ch in 'BCDEFGHIJKL': ws_s.column_dimensions[ch].width = 16
        
        ws_s.merge_cells('B2:J2')
        ws_s['B2'] = "TABCORE - OVERALL SYSTEM INVENTORY DASHBOARD"
        ws_s['B2'].font = Font(bold=True, size=16, color="1A2A6C", name="Calibri")
        ws_s['B2'].alignment = Alignment(horizontal="center")

        ws_s.merge_cells('B3:J3')
        ws_s['B3'] = f"Generated: {display_time}  |  Authorized By: {safe_user_name}"
        ws_s['B3'].font = Font(italic=True, size=9, color="555555", name="Calibri")
        ws_s['B3'].alignment = Alignment(horizontal="center")

        kpis = [
            ("B", "TOTAL DEVICES", stats['total'], "1A2A6C"),
            ("D", "PASSED ✓", stats['passed'], "14532D"),
            ("F", "MINOR ⚠", stats['minor'], "78350F"),
            ("H", "DEFECTIVE ❌", stats['defective'], "7F1D1D"),
            ("J", "DOA ☠", stats['doa'], "4C1D95")
        ]
        
        ws_s.row_dimensions[5].height = 16
        ws_s.row_dimensions[6].height = 40
        
        for col, label, val, color in kpis:
            next_col = get_column_letter(ord(col) - 64 + 1)
            ws_s.merge_cells(f'{col}5:{next_col}5')
            c = ws_s[f'{col}5']
            c.value = label
            c.font = Font(bold=True, size=9, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor=color)
            c.alignment = Alignment(horizontal="center", vertical="center")
            
            ws_s.merge_cells(f'{col}6:{next_col}6')
            c = ws_s[f'{col}6']
            c.value = val
            c.font = Font(bold=True, size=22, color=color)
            c.fill = PatternFill("solid", fgColor="F8FAFF")
            c.border = thin_border()
            c.alignment = Alignment(horizontal="center", vertical="center")

        # Pie Chart Data
        chart_data_start = 13
        chart_items = [("Passed", stats['passed']), ("Minor", stats['minor']), ("Defective", stats['defective']), ("DOA", stats['doa'])]
        for i, (name, val) in enumerate(chart_items):
            ws_s.cell(row=chart_data_start+i, column=25, value=name)
            ws_s.cell(row=chart_data_start+i, column=26, value=val)

        # 🔴 Updated Pie Chart Title (Placed in cells instead of inside the chart to avoid overlapping)
        ws_s.merge_cells('F8:K8')
        title_cell = ws_s['F8']
        title_cell.value = "OVERALL INSPECTION RESULTS"
        title_cell.font = Font(bold=True, size=11, color="1A2A6C")
        title_cell.alignment = Alignment(horizontal="center", vertical="center")

        pie = PieChart()
        pie.title = None # Removed internal overlapping title
        labels = Reference(ws_s, min_col=25, min_row=chart_data_start, max_row=chart_data_start+3)
        data = Reference(ws_s, min_col=26, min_row=chart_data_start-1, max_row=chart_data_start+3)
        pie.add_data(data, titles_from_data=True)
        pie.set_categories(labels)
        pie.height = 10
        pie.width = 16
        
        slices = [DataPoint(idx=i) for i in range(4)]
        slices[0].graphicalProperties.solidFill = "14532D"
        slices[1].graphicalProperties.solidFill = "F0A500"
        slices[2].graphicalProperties.solidFill = "B91C1C"
        slices[3].graphicalProperties.solidFill = "4C1D95"
        pie.series[0].data_points = slices
        
        ws_s.add_chart(pie, "F9")

        # District Breakdown
        dr = 9
        ws_s.merge_cells(f'B{dr}:D{dr}')
        ws_s[f'B{dr}'] = "DISTRICT BREAKDOWN"
        ws_s[f'B{dr}'].font = Font(bold=True, color="FFFFFF")
        ws_s[f'B{dr}'].fill = navy_fill
        ws_s[f'B{dr}'].alignment = Alignment(horizontal="center")
        
        dr += 1
        for ci, h in enumerate(['District', 'Total', 'Passed'], 2):
            c = ws_s.cell(row=dr, column=ci, value=h)
            c.font, c.fill, c.border = Font(bold=True, color="FFFFFF", size=9), navy_fill, thin_border()
            c.alignment = Alignment(horizontal="center")

        dr += 1
        for dn, dc in sorted(district_stats.items()):
            for ci, v in enumerate([dn, dc['total'], dc['passed']], 2):
                c = ws_s.cell(row=dr, column=ci, value=v)
                c.fill = alt_fill if dr % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
                c.border, c.alignment = thin_border(), Alignment(horizontal="left" if ci==2 else "center")
            dr += 1

        # =========================================================
        # SHEET 2: FULL INVENTORY
        # =========================================================
        ws_d = wb.create_sheet("📋 Full Inventory")
        ws_d.sheet_properties.tabColor = "14532D"
        ws_d.sheet_view.showGridLines = False
        ws_d.freeze_panes = "A5"

        ws_d.merge_cells('A2:M2')
        ws_d['A2'] = f"TABCORE MASTER DATABASE EXPORT  |  {display_time}"
        ws_d['A2'].font = Font(bold=True, size=14, color="1A2A6C")

        base_headers = ['T-ID','District','Brand','Model','Serial Number','Asset No','IMEI Number','Status','Registered By','Inspected By']
        all_headers = base_headers + ITEM_LABELS + ['Notes', 'Charger', 'Cable', 'SIM Pin']

        ws_d.append([])
        ws_d.append(all_headers)
        ws_d.row_dimensions[4].height = 20

        for ci, h in enumerate(all_headers, 1):
            c = ws_d.cell(row=4, column=ci, value=h)
            c.font, c.fill, c.border = Font(bold=True, color="FFFFFF", size=9), navy_fill, thin_border()
            c.alignment = Alignment(horizontal="center", vertical="center")

        ws_d.auto_filter.ref = f"A4:{get_column_letter(len(all_headers))}4"

        for ri, t in enumerate(tablets, 5):
            st = str(t.get('status',''))
            rf = STATUS_FILL.get(st, PatternFill("solid", fgColor="FFFFFF"))
            
            details = {}
            if t.get('inspection_data'):
                try: details = json.loads(t['inspection_data'])
                except: pass

            row_data = [
                f"T-{t['id']}", str(t.get('district','-')).title(), t.get('brand','-'), t.get('model','-'),
                str(t.get('serial_number','-')), t.get('asset_no','-'), str(t.get('imei_number','-')),
                st, t.get('registered_by','-'), t.get('inspected_by','-')
            ]
            
            for key in item_keys:
                v = str(details.get(key,'-')).lower()
                row_data.append('Pass' if v=='pass' else 'Minor' if v in ['minor','partial'] else 'Fail' if v=='fail' else 'N/A' if v=='n/a' else '-')
                
            row_data.append(details.get('inspector_notes','-') or '-')
            row_data.extend([t.get('charger_status','-'), t.get('cable_status','-'), t.get('simpin_status','-')])

            for ci, val in enumerate(row_data, 1):
                c = ws_d.cell(row=ri, column=ci, value=val)
                c.font, c.border = Font(size=9, bold=(ci==8)), thin_border()
                c.fill = rf if ci == 8 else (alt_fill if ri % 2 == 0 else PatternFill("solid", fgColor="FFFFFF"))
                c.alignment = Alignment(vertical="center", horizontal="left" if ci in [5,7] else "center")

        for col_idx in range(1, len(all_headers) + 1):
            max_len = 0
            for row in ws_d.iter_rows(min_row=4, min_col=col_idx, max_col=col_idx):
                for cell in row:
                    try:
                        if cell.value and len(str(cell.value)) > max_len: max_len = len(str(cell.value))
                    except: pass
            ws_d.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 2, 35)

        # =========================================================
        # SHEET 3: ACCESSORY SHORTAGES
        # =========================================================
        if shortages_list:
            ws_a = wb.create_sheet("🔌 Accessory Shortages")
            ws_a.sheet_properties.tabColor = "B91C1C"
            ws_a.sheet_view.showGridLines = False
            
            ws_a.merge_cells('A2:H2')
            # 🔴 Updated Title 
            ws_a['A2'] = f"ACCESSORY SHORTAGES REPORT (Missing & Damaged Items)  |  Total Devices: {len(shortages_list)}"
            ws_a['A2'].font = Font(bold=True, size=12, color="B91C1C")

            a_hdrs = ['T-ID', 'District', 'Serial Number', 'Device Status', 'Charger', 'Cable', 'SIM Pin']
            ws_a.append([])
            ws_a.append(a_hdrs)
            
            for ci, h in enumerate(a_hdrs, 1):
                c = ws_a.cell(row=4, column=ci, value=h)
                c.font, c.fill, c.border = Font(bold=True, color="FFFFFF", size=9), PatternFill("solid", fgColor="7F1D1D"), thin_border()
                c.alignment = Alignment(horizontal="center")

            for ri, s in enumerate(shortages_list, 5):
                row_data = [f"T-{s['id']}", s['district'], s['serial'], s['status'], s['charger'], s['cable'], s['sim_pin']]
                for ci, val in enumerate(row_data, 1):
                    c = ws_a.cell(row=ri, column=ci, value=val)
                    
                    # 🔴 Updated formatting logic for both Missing & Damaged
                    val_lower = str(val).lower()
                    is_issue = 'missing' in val_lower or 'damage' in val_lower
                    
                    c.font, c.border = Font(size=9, bold=is_issue), thin_border()
                    c.fill = PatternFill("solid", fgColor="F8D7DA") if is_issue else PatternFill("solid", fgColor="FFFFFF")
                    c.alignment = Alignment(horizontal="center")
            
            for i, w in enumerate([8, 15, 20, 15, 15, 15, 15], 1): ws_a.column_dimensions[get_column_letter(i)].width = w

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        filename = f"Tabcore_Master_Inventory_{datetime.now(tz).strftime('%Y%m%d_%H%M')}.xlsx"
        
        try: log_system_audit("Full Export", safe_user_name, "Exported Complete Database with Master Dashboard.")
        except: pass

        return Response(output, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment;filename={filename}"})
        
    except Exception as e:
        import traceback
        return f"Error: {str(e)}<br><pre>{traceback.format_exc()}</pre>"
        
@app.route('/export_filtered', methods=['GET'])
@login_required
def export_filtered():
    try:
        safe_user_name = "System Admin"
        try:
            if current_user:
                if hasattr(current_user, 'name') and current_user.name:
                    safe_user_name = current_user.name
                elif hasattr(current_user, 'username') and current_user.username:
                    safe_user_name = current_user.username
        except: pass

        if hasattr(current_user, 'role') and current_user.role != 'Admin':
            return redirect(url_for('dashboard'))

        from openpyxl import Workbook
        from openpyxl.styles import (Font, Alignment, PatternFill,
                                     Border, Side, GradientFill)
        from openpyxl.utils import get_column_letter
        from openpyxl.chart import BarChart, PieChart, Reference
        from openpyxl.chart.series import DataPoint
        from datetime import datetime
        import pytz, io, json

        conn = get_db_connection()
        cursor = conn.cursor()

        district = request.args.get('district', '')
        brand    = request.args.get('brand', '')
        status   = request.args.get('status', '')
        inspector= request.args.get('inspector', '')

        query = ("SELECT * FROM tablets "
                 "WHERE status NOT IN ('Pending','Locked') "
                 "AND is_deleted = 0")
        params = []
        filter_texts  = []
        filename_parts= ["Tabcore", "Export"]

        if district and district != 'All Districts':
            query += " AND district = %s"; params.append(district)
            filter_texts.append(f"District: {district}")
            filename_parts.append(district)
        if brand and brand != 'All Brands':
            query += " AND brand = %s"; params.append(brand)
            filter_texts.append(f"Brand: {brand}")
            filename_parts.append(brand)
        if status and status != 'All Statuses':
            query += " AND status = %s"; params.append(status)
            filter_texts.append(f"Status: {status}")
            filename_parts.append(status)
        if inspector and inspector != 'All Inspectors':
            query += " AND inspected_by = %s"; params.append(inspector)
            filter_texts.append(f"Inspector: {inspector}")
            filename_parts.append(inspector.replace(" ", ""))

        query += " ORDER BY district ASC, id ASC"
        cursor.execute(query, params)
        tablets = cursor.fetchall()
        conn.close()

        if not tablets:
            return "No data found for the selected filters."

        tz           = pytz.timezone('Asia/Colombo')
        now          = datetime.now(tz)
        display_time = now.strftime('%Y-%m-%d  %I:%M %p')
        filter_str   = ', '.join(filter_texts) if filter_texts else 'None (All Data)'

        # ── Shared styles ──────────────────────────────────────────────
        def thin_border():
            s = Side(style='thin', color='CCCCCC')
            return Border(left=s, right=s, top=s, bottom=s)

        def thick_border():
            s = Side(style='medium', color='1A2A6C')
            return Border(left=s, right=s, top=s, bottom=s)

        navy_fill   = PatternFill("solid", fgColor="1A2A6C")
        red_fill    = PatternFill("solid", fgColor="7F1D1D")
        green_fill  = PatternFill("solid", fgColor="14532D")
        amber_fill  = PatternFill("solid", fgColor="78350F")
        purple_fill = PatternFill("solid", fgColor="4C1D95")

        pass_fill   = PatternFill("solid", fgColor="D4EDDA")
        minor_fill  = PatternFill("solid", fgColor="FFF3CD")
        fail_fill   = PatternFill("solid", fgColor="F8D7DA")
        doa_fill    = PatternFill("solid", fgColor="EDE9FE")
        alt_fill    = PatternFill("solid", fgColor="F8FAFF")

        STATUS_FILL = {
            'Passed':            pass_fill,
            'Minor Issues':      minor_fill,
            'Defective':         fail_fill,
            'Dead Device (DOA)': doa_fill,
        }

        def hdr_cell(ws, row, col, value, fill=None, font_color="FFFFFF",
                     size=10, bold=True, align="center", wrap=False):
            c = ws.cell(row=row, column=col, value=value)
            c.font      = Font(bold=bold, color=font_color, size=size,
                               name="Calibri")
            c.fill      = fill or navy_fill
            c.alignment = Alignment(horizontal=align, vertical="center",
                                    wrap_text=wrap)
            c.border    = thin_border()
            return c

        def data_cell(ws, row, col, value, fill=None, bold=False,
                      font_color="2D2D2D", align="center"):
            c = ws.cell(row=row, column=col, value=value)
            c.font      = Font(bold=bold, color=font_color, size=9,
                               name="Calibri")
            c.fill      = fill or PatternFill("solid", fgColor="FFFFFF")
            c.alignment = Alignment(horizontal=align, vertical="center",
                                    wrap_text=True)
            c.border    = thin_border()
            return c

        # ── Pre-compute stats ──────────────────────────────────────────
        item_keys = ['display','touch','battery','cameras','wifi','bt',
                     'gps','speaker','mic','charging','p_btn','sim']
        ITEM_LABELS = ['Display','Touch','Battery','Camera','Wi-Fi',
                       'Bluetooth','GPS','Speaker','Mic',
                       'Charging Port','Power Btn','SIM Slot']

        stats = {'total':len(tablets),'passed':0,'minor':0,
                 'defective':0,'doa':0}
        district_stats  = {}
        inspector_stats = {}

        miss_chg = dmg_chg = miss_cbl = dmg_cbl = miss_pin = 0

        for t in tablets:
            st = t.get('status','')
            if   st == 'Passed':             stats['passed']    += 1
            elif st == 'Minor Issues':        stats['minor']     += 1
            elif st == 'Defective':           stats['defective'] += 1
            elif st == 'Dead Device (DOA)':   stats['doa']       += 1

            d = str(t.get('district') or 'Unknown').strip().title()
            if d not in district_stats:
                district_stats[d] = {'total':0,'passed':0,'minor':0,
                                     'defective':0,'doa':0}
            district_stats[d]['total'] += 1
            if   st == 'Passed':             district_stats[d]['passed']    += 1
            elif st == 'Minor Issues':        district_stats[d]['minor']     += 1
            elif st == 'Defective':           district_stats[d]['defective'] += 1
            elif st == 'Dead Device (DOA)':   district_stats[d]['doa']       += 1

            insp = str(t.get('inspected_by') or 'Unknown')
            if insp not in inspector_stats:
                inspector_stats[insp] = {'total':0,'passed':0,'minor':0,
                                         'defective':0,'doa':0}
            inspector_stats[insp]['total'] += 1
            if   st == 'Passed':             inspector_stats[insp]['passed']    += 1
            elif st == 'Minor Issues':        inspector_stats[insp]['minor']     += 1
            elif st == 'Defective':           inspector_stats[insp]['defective'] += 1
            elif st == 'Dead Device (DOA)':   inspector_stats[insp]['doa']       += 1

            cs = str(t.get('charger_status','')).lower()
            cb = str(t.get('cable_status','')).lower()
            sp = str(t.get('simpin_status','')).lower()
            if 'missing' in cs: miss_chg += 1
            elif 'damage' in cs: dmg_chg += 1
            if 'missing' in cb: miss_cbl += 1
            elif 'damage' in cb: dmg_cbl += 1
            if 'missing' in sp: miss_pin += 1

        wb = Workbook()

        # ══════════════════════════════════════════════════════════════
        # SHEET 1 — SUMMARY DASHBOARD
        # ══════════════════════════════════════════════════════════════
        ws_s = wb.active
        ws_s.title = "📊 Summary"
        ws_s.sheet_properties.tabColor = "1A2A6C"
        ws_s.sheet_view.showGridLines = False

        # Title block
        ws_s.row_dimensions[1].height = 8
        ws_s.row_dimensions[2].height = 36
        ws_s.row_dimensions[3].height = 20
        ws_s.row_dimensions[4].height = 18

        ws_s.column_dimensions['A'].width = 3
        for ch in 'BCDEFGHIJ':
            ws_s.column_dimensions[ch].width = 18

        ws_s.merge_cells('B2:J2')
        c = ws_s['B2']
        c.value     = "DEPARTMENT OF CENSUS AND STATISTICS — TABCORE SYSTEM"
        c.font      = Font(bold=True, size=16, color="1A2A6C", name="Calibri")
        c.alignment = Alignment(horizontal="center", vertical="center")

        ws_s.merge_cells('B3:J3')
        c = ws_s['B3']
        c.value     = f"Inspection Summary Dashboard  |  {display_time}  |  By: {safe_user_name}"
        c.font      = Font(italic=True, size=9, color="555555", name="Calibri")
        c.alignment = Alignment(horizontal="center")

        ws_s.merge_cells('B4:J4')
        c = ws_s['B4']
        c.value     = f"Filters: {filter_str}"
        c.font      = Font(italic=True, size=9, color="7F1D1D", name="Calibri")
        c.alignment = Alignment(horizontal="center")

        # KPI Cards row
        ws_s.row_dimensions[6].height = 14
        ws_s.row_dimensions[7].height = 44
        ws_s.row_dimensions[8].height = 20

        kpis = [
            ("B","C", "TOTAL",    stats['total'],    navy_fill,   "FFFFFF"),
            ("D","E", "PASSED ✓", stats['passed'],   green_fill,  "FFFFFF"),
            ("F","F", "MINOR ⚠",  stats['minor'],    amber_fill,  "FFFFFF"),
            ("G","H", "DEFECTIVE",stats['defective'],red_fill,    "FFFFFF"),
            ("I","J", "DOA ☠",   stats['doa'],      purple_fill, "FFFFFF"),
        ]
        for start_col, end_col, label, val, bg, fc in kpis:
            ws_s.merge_cells(f'{start_col}6:{end_col}6')
            c = ws_s[f'{start_col}6']
            c.value     = label
            c.font      = Font(bold=True, size=9, color=fc, name="Calibri")
            c.fill      = bg
            c.alignment = Alignment(horizontal="center", vertical="center")

            ws_s.merge_cells(f'{start_col}7:{end_col}7')
            c = ws_s[f'{start_col}7']
            c.value     = val
            c.font      = Font(bold=True, size=24, color=fc, name="Calibri")
            c.fill      = bg
            c.alignment = Alignment(horizontal="center", vertical="center")

        # District breakdown table
        dr = 10
        ws_s.merge_cells(f'B{dr}:J{dr}')
        c = ws_s[f'B{dr}']
        c.value     = "DISTRICT BREAKDOWN"
        c.font      = Font(bold=True, size=11, color="FFFFFF", name="Calibri")
        c.fill      = navy_fill
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws_s.row_dimensions[dr].height = 22

        dr += 1
        d_hdrs = ['District','Total','Passed','Minor','Defective','DOA','Pass Rate']
        d_cols = 'BCDEFGH'
        for ci, h in enumerate(d_hdrs):
            hdr_cell(ws_s, dr, ord(d_cols[ci])-64, h)
        ws_s.row_dimensions[dr].height = 18

        dr += 1
        for dn, dc in sorted(district_stats.items()):
            rate = (f"{round(dc['passed']/dc['total']*100)}%"
                    if dc['total'] > 0 else "0%")
            vals = [dn, dc['total'], dc['passed'],
                    dc['minor'], dc['defective'], dc['doa'], rate]
            for ci, v in enumerate(vals):
                f = alt_fill if dr % 2 == 0 else PatternFill("solid",fgColor="FFFFFF")
                data_cell(ws_s, dr, ord(d_cols[ci])-64, v, fill=f,
                          align="left" if ci==0 else "center")
            ws_s.row_dimensions[dr].height = 16
            dr += 1

        # Accessory exception summary
        dr += 1
        ws_s.merge_cells(f'B{dr}:J{dr}')
        c = ws_s[f'B{dr}']
        c.value     = "ACCESSORY EXCEPTION SUMMARY"
        c.font      = Font(bold=True, size=11, color="FFFFFF", name="Calibri")
        c.fill      = red_fill
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws_s.row_dimensions[dr].height = 22

        dr += 1
        acc_hdrs = ['Item','Missing','Damaged','Total Issues']
        for ci, h in enumerate(acc_hdrs):
            hdr_cell(ws_s, dr, ci+2, h, fill=red_fill)
        ws_s.row_dimensions[dr].height = 18

        acc_data = [
            ('Charger',  miss_chg, dmg_chg),
            ('Data Cable',miss_cbl, dmg_cbl),
            ('SIM Pin',  miss_pin, 0),
        ]
        dr += 1
        for item, miss, dmg in acc_data:
            total_issue = miss + dmg
            f = PatternFill("solid", fgColor="FFF5F5") if total_issue > 0 else PatternFill("solid", fgColor="F0FAF4")
            data_cell(ws_s, dr, 2, item,        fill=f, align="left", bold=True)
            data_cell(ws_s, dr, 3, miss,        fill=f, font_color="B91C1C" if miss > 0 else "2D2D2D")
            data_cell(ws_s, dr, 4, dmg,         fill=f, font_color="B45309" if dmg > 0 else "2D2D2D")
            data_cell(ws_s, dr, 5, total_issue, fill=f, bold=True)
            ws_s.row_dimensions[dr].height = 16
            dr += 1

        # ── Bar Chart ─────────────────────────────────────────────────
        if len(district_stats) > 0:
            chart_start_row = 6
            chart_data_row  = 11

            bar = BarChart()
            bar.type    = "col"
            bar.title   = "Inspection Results by District"
            bar.style   = 10
            bar.width   = 22
            bar.height  = 12
            bar.grouping = "clustered"

            cats = Reference(ws_s,
                             min_col=2,
                             min_row=chart_data_row + 1,
                             max_row=chart_data_row + len(district_stats))

            for col_offset, series_name in [(4, "Passed"),
                                             (5, "Minor"),
                                             (6, "Defective")]:
                data = Reference(ws_s,
                                 min_col=col_offset,
                                 min_row=chart_data_row,
                                 max_row=chart_data_row + len(district_stats))
                bar.add_data(data, titles_from_data=True)

            bar.set_categories(cats)
            bar.series[0].graphicalProperties.solidFill = "14532D"
            bar.series[1].graphicalProperties.solidFill = "F0A500"
            if len(bar.series) > 2:
                bar.series[2].graphicalProperties.solidFill = "B91C1C"

            ws_s.add_chart(bar, "B" + str(dr + 2))

        # ══════════════════════════════════════════════════════════════
        # SHEET 2 — INSPECTION DATA (Full)
        # ══════════════════════════════════════════════════════════════
        ws_d = wb.create_sheet("📋 Inspection Data")
        ws_d.sheet_properties.tabColor = "1A2A6C"
        ws_d.sheet_view.showGridLines  = False
        ws_d.freeze_panes = "A6"

        ws_d.row_dimensions[1].height = 8
        ws_d.row_dimensions[2].height = 32
        ws_d.row_dimensions[3].height = 18
        ws_d.row_dimensions[4].height = 16

        ws_d.merge_cells('A2:Z2')
        c = ws_d['A2']
        c.value     = "TABCORE — FULL INSPECTION DATA"
        c.font      = Font(bold=True, size=14, color="1A2A6C", name="Calibri")
        c.alignment = Alignment(horizontal="center", vertical="center")

        ws_d.merge_cells('A3:Z3')
        c = ws_d['A3']
        c.value     = (f"Generated: {display_time}  |  By: {safe_user_name}"
                       f"  |  Filters: {filter_str}")
        c.font      = Font(italic=True, size=8, color="555555", name="Calibri")
        c.alignment = Alignment(horizontal="center")

        base_headers = ['T-ID','District','Brand','Model','Serial Number',
                        'Asset No','IMEI Number','Status',
                        'Registered By','Inspected By','Battery Drain',
                        'Charger','Cable','SIM Pin']
        hw_headers   = ITEM_LABELS + ['Inspector Notes']
        all_headers  = base_headers + hw_headers

        ws_d.append([])
        ws_d.append(all_headers)
        ws_d.row_dimensions[5].height = 22

        for ci, h in enumerate(all_headers, 1):
            c = ws_d.cell(row=5, column=ci, value=h)
            c.font      = Font(bold=True, color="FFFFFF", size=8.5, name="Calibri")
            c.fill      = navy_fill
            c.alignment = Alignment(horizontal="center", vertical="center",
                                    wrap_text=True)
            c.border    = thin_border()

        last_col = get_column_letter(len(all_headers))
        ws_d.auto_filter.ref = f"A5:{last_col}5"

        for ri, t in enumerate(tablets, 6):
            st = t.get('status','')
            rf = STATUS_FILL.get(st, PatternFill("solid", fgColor="FFFFFF"))
            af = alt_fill if ri % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")

            details = {}
            if t.get('inspection_data'):
                try: details = json.loads(t['inspection_data'])
                except: pass

            base_row = [
                f"T-{t['id']}",
                str(t.get('district','-')).title(),
                t.get('brand','-'),
                t.get('model','-'),
                str(t.get('serial_number','-')),
                t.get('asset_no','-'),
                str(t.get('imei_number','-')),
                st,
                t.get('registered_by','-'),
                t.get('inspected_by','-') or '-',
                t.get('battery_drain_time','-') or '-',
                t.get('charger_status','-'),
                t.get('cable_status','-'),
                t.get('simpin_status','-'),
            ]
            hw_row = []
            for key in item_keys:
                v = str(details.get(key,'-')).lower()
                hw_row.append(
                    'Pass'   if v == 'pass'                  else
                    'Minor'  if v in ['minor','partial']      else
                    'Fail'   if v == 'fail'                   else
                    'N/A'    if v == 'n/a'                    else '-'
                )
            hw_row.append(details.get('inspector_notes','-') or '-')

            full_row = base_row + hw_row
            for ci, val in enumerate(full_row, 1):
                c = ws_d.cell(row=ri, column=ci, value=val)
                c.font      = Font(size=8.5, name="Calibri")
                c.alignment = Alignment(vertical="center",
                                        horizontal="left" if ci in [5,7,15] else "center",
                                        wrap_text=(ci == len(full_row)))
                c.border    = thin_border()

                # Status cell special color
                if ci == 8:
                    c.fill = rf
                    c.font = Font(bold=True, size=8.5, name="Calibri",
                                  color=("155724" if st=='Passed' else
                                         "856404" if st=='Minor Issues' else
                                         "721C24" if st=='Defective' else
                                         "4C1D95"))
                # HW fail cells
                elif ci >= 15 and ci < 15 + len(item_keys):
                    v = str(val).lower()
                    c.fill = (PatternFill("solid",fgColor="D4EDDA") if v=='pass'  else
                              PatternFill("solid",fgColor="FFF3CD") if v=='minor' else
                              PatternFill("solid",fgColor="F8D7DA") if v=='fail'  else
                              af)
                else:
                    c.fill = af

            ws_d.row_dimensions[ri].height = 16

        # Column widths
        col_widths = [8,12,10,12,18,10,18,14,14,14,12,10,10,10]
        for ci, w in enumerate(col_widths, 1):
            ws_d.column_dimensions[get_column_letter(ci)].width = w
        for ci in range(len(col_widths)+1, len(all_headers)+1):
            ws_d.column_dimensions[get_column_letter(ci)].width = 8
        # Notes column wide
        ws_d.column_dimensions[get_column_letter(len(all_headers))].width = 35

        # Print setup
        ws_d.page_setup.orientation = 'landscape'
        ws_d.page_setup.fitToWidth  = 1
        ws_d.page_setup.fitToHeight = 0
        ws_d.print_title_rows = '5:5'

        # ══════════════════════════════════════════════════════════════
        # SHEET 3 — DEFECTS ONLY
        # ══════════════════════════════════════════════════════════════
        defect_tablets = [t for t in tablets
                          if t.get('status') in
                          ('Minor Issues','Defective','Dead Device (DOA)')]

        if defect_tablets:
            ws_def = wb.create_sheet("⚠ Defects")
            ws_def.sheet_properties.tabColor = "7F1D1D"
            ws_def.sheet_view.showGridLines  = False
            ws_def.freeze_panes = "A5"

            ws_def.row_dimensions[1].height = 8
            ws_def.row_dimensions[2].height = 30

            ws_def.merge_cells('A2:L2')
            c = ws_def['A2']
            c.value     = f"DEFECTIVE & FLAGGED DEVICES  |  Total: {len(defect_tablets)}  |  {display_time}"
            c.font      = Font(bold=True, size=12, color="7F1D1D", name="Calibri")
            c.alignment = Alignment(horizontal="center", vertical="center")

            def_hdrs = ['No.','T-ID','Asset No','District','Serial Number',
                        'Status','Action Required',
                        'Hardware Faults','Accessory Issues','Inspector Notes']
            ws_def.append([])
            ws_def.append(def_hdrs)
            ws_def.row_dimensions[4].height = 20

            for ci, h in enumerate(def_hdrs, 1):
                c = ws_def.cell(row=4, column=ci, value=h)
                c.font      = Font(bold=True, color="FFFFFF", size=9, name="Calibri")
                c.fill      = red_fill
                c.alignment = Alignment(horizontal="center", vertical="center")
                c.border    = thin_border()

            # 🔴 මෙන්න මෙතන තමයි වෙනස් කළේ! Supplier Return අයින් කරලා Send to Repair දැම්මා.
            ACTION_MAP = {
                'Minor Issues':      'Monitor / Note',
                'Defective':         'Send to Repair',
                'Dead Device (DOA)': 'Send to Repair',
            }
            HW_MAP = {k:l for k,l in zip(item_keys, ITEM_LABELS)}

            for ri, t in enumerate(defect_tablets, 5):
                st  = t.get('status','')
                rf  = STATUS_FILL.get(st, PatternFill("solid",fgColor="FFFFFF"))
                act = ACTION_MAP.get(st, '-')

                details = {}
                if t.get('inspection_data'):
                    try: details = json.loads(t['inspection_data'])
                    except: pass

                # HW faults
                hw_faults = []
                for k, lbl in HW_MAP.items():
                    v = str(details.get(k,'')).lower()
                    if   v == 'fail':                hw_faults.append(f"{lbl}: FAIL")
                    elif v in ['minor','partial']:    hw_faults.append(f"{lbl}: Minor")
                hw_str = " | ".join(hw_faults) if hw_faults else "-"

                # Acc issues
                acc_issues = []
                for key, lbl in [('charger_status','Charger'),
                                  ('cable_status','Cable'),
                                  ('simpin_status','SIM Pin')]:
                    v = str(t.get(key,'')).lower()
                    if   'missing' in v: acc_issues.append(f"{lbl}: Missing")
                    elif 'damage'  in v: acc_issues.append(f"{lbl}: Damaged")
                acc_str = " | ".join(acc_issues) if acc_issues else "Full Set"

                notes = str(details.get('inspector_notes','') or '-')

                row_data = [
                    ri - 4,
                    f"T-{t['id']}",
                    t.get('asset_no','-'),
                    str(t.get('district','-')).title(),
                    t['serial_number'],
                    st, act, hw_str, acc_str, notes
                ]
                col_widths_def = [6,8,10,12,18,14,16,40,25,40]

                for ci, val in enumerate(row_data, 1):
                    c = ws_def.cell(row=ri, column=ci, value=val)
                    c.font      = Font(size=8.5, name="Calibri",
                                       bold=(ci in (6,7)))
                    c.fill      = rf
                    c.alignment = Alignment(vertical="center",
                                            horizontal="left" if ci > 5 else "center",
                                            wrap_text=True)
                    c.border    = thin_border()

                ws_def.row_dimensions[ri].height = 30

            for ci, w in enumerate(col_widths_def, 1):
                ws_def.column_dimensions[get_column_letter(ci)].width = w

            ws_def.page_setup.orientation = 'landscape'
            ws_def.page_setup.fitToWidth  = 1
            ws_def.print_title_rows = '4:4'

        # ══════════════════════════════════════════════════════════════
        # SHEET 4 — INSPECTOR PERFORMANCE
        # ══════════════════════════════════════════════════════════════
        if inspector_stats:
            ws_i = wb.create_sheet("👤 Inspector Stats")
            ws_i.sheet_properties.tabColor = "14532D"
            ws_i.sheet_view.showGridLines  = False

            ws_i.row_dimensions[1].height = 8
            ws_i.row_dimensions[2].height = 30
            ws_i.row_dimensions[3].height = 16

            ws_i.merge_cells('A2:H2')
            c = ws_i['A2']
            c.value     = f"INSPECTOR PERFORMANCE SUMMARY  |  {display_time}"
            c.font      = Font(bold=True, size=12, color="14532D", name="Calibri")
            c.alignment = Alignment(horizontal="center", vertical="center")

            i_hdrs = ['Inspector','Total','Passed','Minor',
                      'Defective','DOA','Pass Rate','Defect Rate']
            ws_i.append([])
            ws_i.append(i_hdrs)
            ws_i.row_dimensions[4].height = 20

            for ci, h in enumerate(i_hdrs, 1):
                c = ws_i.cell(row=4, column=ci, value=h)
                c.font      = Font(bold=True, color="FFFFFF", size=9,
                                   name="Calibri")
                c.fill      = green_fill
                c.alignment = Alignment(horizontal="center", vertical="center")
                c.border    = thin_border()

            for ri, (insp, ist) in enumerate(
                    sorted(inspector_stats.items(),
                           key=lambda x: x[1]['total'], reverse=True),
                    start=5):
                t_  = ist['total']
                p_r = f"{round(ist['passed']/t_*100)}%" if t_ > 0 else "0%"
                d_r = f"{round((ist['defective']+ist['doa'])/t_*100)}%" if t_ > 0 else "0%"
                f   = alt_fill if ri % 2 == 0 else PatternFill("solid",fgColor="FFFFFF")

                row_data = [insp, t_, ist['passed'], ist['minor'],
                            ist['defective'], ist['doa'], p_r, d_r]
                for ci, val in enumerate(row_data, 1):
                    c = ws_i.cell(row=ri, column=ci, value=val)
                    c.font      = Font(size=9, name="Calibri",
                                       bold=(ci == 1))
                    c.fill      = f
                    c.alignment = Alignment(
                        horizontal="left" if ci==1 else "center",
                        vertical="center")
                    c.border    = thin_border()
                ws_i.row_dimensions[ri].height = 16

            col_ws_i = [22, 10, 10, 10, 12, 10, 12, 12]
            for ci, w in enumerate(col_ws_i, 1):
                ws_i.column_dimensions[get_column_letter(ci)].width = w

        # ── Save & return ──────────────────────────────────────────────
        # Move Summary to first position
        wb.move_sheet("📊 Summary", offset=-wb.index(wb["📊 Summary"]))

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        filename = "_".join(filename_parts) + ".xlsx"

        try:
            log_system_audit("Excel Export", safe_user_name,
                             f"Exported Professional Excel. Filters: {filter_str}")
        except: pass

        return Response(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition":
                     f"attachment;filename={filename}"})

    except Exception as e:
        import traceback
        return (f"Error generating Excel: {str(e)}"
                f"<br><pre>{traceback.format_exc()}</pre>")
        
@app.route('/export_sinhala_word')
def export_sinhala_word():
    target_district = request.args.get('district', 'All Districts')
    target_brand = request.args.get('brand', 'All Brands')
    target_status = request.args.get('status', 'All Statuses')

    conn = get_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    query = "SELECT * FROM tablets WHERE 1=1"
    params = []
    
    if target_district != 'All Districts':
        query += " AND district = %s"
        params.append(target_district)
    if target_brand != 'All Brands':
        query += " AND brand = %s"
        params.append(target_brand)
    if target_status != 'All Statuses':
        query += " AND status = %s"
        params.append(target_status)
        
    cursor.execute(query, tuple(params))
    tablets = cursor.fetchall()
    conn.close()

    defective_tablets = []
    good_count = 0
    for t in tablets:
        issues = []
        
        if t.get('status') in ['Defective', 'Fail']:
            issues.append("දෘඩාංග දෝෂ සහිතයි (අලුත්වැඩියා කළ යුතුය)")
        elif t.get('status') == 'Dead Device (DOA)':
            issues.append("උපාංගය සම්පූර්ණයෙන්ම අක්‍රියයි (Dead Device / DOA)")
        
        det = t.get('details', {})
        if isinstance(det, str): 
            import json
            try: det = json.loads(det)
            except: det = {}

        if det.get('battery') == 'fail': issues.append("බැටරි දෝෂ සහිතයි")
        if det.get('display') == 'fail': issues.append("තිරය (Screen) දෝෂ සහිතයි")

        if issues:
            defective_tablets.append({
                'asset': t.get('asset_no', '-'), 
                'sn': t.get('serial_number', '-'), 
                'issue': " | ".join(issues)
            })
        else:
            good_count += 1

    doc = Document()
    doc.add_heading('දිස්ත්‍රික්ක අතර හුවමාරු කරන ලද ටැබ්ලට් පරිගණක වල තත්ත්ව පරීක්ෂණ වාර්තාව', level=1)
    doc.add_paragraph(f"දිස්ත්‍රික්කය: {target_district}")
    doc.add_paragraph(f"පරීක්ෂා කරන ලද මුළු ටැබ්ලට් සංඛ්‍යාව: {len(tablets)}")
    doc.add_paragraph("\nපහත වගුවේ දැක්වෙන ටැබ්ලට් පරිගණක වල දෝෂ පවතින බව නිරීක්ෂණය විය:")

    if defective_tablets:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Asset No'
        hdr_cells[1].text = 'Serial Number'
        hdr_cells[2].text = 'පවතින දෝෂය'
        
        for dt in defective_tablets:
            row_cells = table.add_row().cells
            row_cells[0].text = str(dt['asset'])
            row_cells[1].text = str(dt['sn'])
            row_cells[2].text = dt['issue']
    else:
        doc.add_paragraph("දෝෂ සහිත ටැබ්ලට් කිසිවක් වාර්තා වී නොමැත.")

    doc.add_paragraph(f"\nඒ අනුව, ඉහත දෝෂ සහිත ටැබ්ලට් හැර ඉතිරි ටැබ්ලට් පරිගණක {good_count} ක ප්‍රමාණය ක්ෂේත්‍ර කටයුතු සඳහා යොදාගත හැකි බව නිර්දේශ කරමි.")
    doc.add_paragraph("\n........................\nපරීක්ෂා කළ නිලධාරියා")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(output, as_attachment=True, download_name='Quality_Inspection_Report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
@app.route('/export_single/<int:id>')
@login_required
def export_single(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM tablets WHERE id = %s", (id,))
    t = cursor.fetchone()
    conn.close()
    if not t: return redirect(url_for('dashboard'))
    
    output = io.StringIO()
    writer = csv.writer(output)
    base_headers = ['ID', 'Dist', 'Brand', 'Model', 'S/N', 'Asset No', 'IMEI', 'Status', 'Reg.By', 'Insp.By']
    item_names = ['Display', 'Touch', 'Battery', 'Cam', 'WiFi', 'BT', 'GPS', 'Speaker', 'Mic', 'Port', 'Pwr', 'SIM', 'Notes']
    item_keys = ['display', 'touch', 'battery', 'cameras', 'wifi', 'bt', 'gps', 'speaker', 'mic', 'charging', 'p_btn', 'sim']
    writer.writerow(base_headers + item_names)
    
    safe_sn = safe_excel_text(t.get('serial_number'))
    safe_imei = safe_excel_text(t.get('imei_number'))

    row = [f"T-{t['id']}", t.get('district','-'), t.get('brand','-'), t.get('model','-'), safe_sn, t.get('asset_no', '-'), safe_imei, t.get('status','-'), t.get('registered_by','-'), t.get('inspected_by','-'), t.get('battery_drain_time', '-')]
    details = {}
    if t.get('inspection_data'):
        try: details = json.loads(t['inspection_data'])
        except: pass
        
    for key in item_keys:
        raw_val = details.get(key, '-')
        val = str(raw_val).lower() if raw_val != '-' else '-'
        if val == 'pass': val = 'P'
        elif val == 'minor' or val == 'partial': val = 'M'
        elif val == 'fail': val = 'F'
        elif val == 'n/a': val = 'N/A'
        else: val = raw_val
        row.append(val)
        
    row.append(details.get('inspector_notes', '-'))
    writer.writerow(row)
    return Response(output.getvalue(), mimetype="text/csv", headers={"Content-Disposition": f"attachment;filename=Tabcore_T-{id}_Report.csv"})

@app.route('/download_template')
@login_required
def download_template():
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Bulk_Upload"
        headers = ['District', 'Brand', 'Model', 'Serial Number', 'IMEI Number', 'Asset No', 'Charger Status', 'Cable Status', 'Simpin Status', 'Pouch Status', 'Pen Status', 'Doc Status']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True)
            
        for i in range(1, 1000):
            ws[f"E{i}"].number_format = '@'

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return Response(output, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment;filename=Tabcore_Bulk_Template.xlsx"})
    except:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['District', 'Brand', 'Model', 'Serial Number', 'IMEI Number', 'Asset No', 'Charger Status', 'Cable Status', 'Simpin Status', 'Pouch Status', 'Pen Status', 'Doc Status'])
        return Response(output.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment;filename=Tabcore_Bulk_Template.csv"})

# --- ADMIN ROUTES & OTHERS ---

@app.route('/officers', methods=['GET', 'POST'])
@login_required
def officers():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    conn = get_db_connection()
    cursor = conn.cursor()
    if request.method == 'POST':
        hashed_pw = generate_password_hash(request.form.get('password'))
        try:
            cursor.execute("INSERT INTO users (username, password, name, role) VALUES (%s, %s, %s, %s)", (request.form.get('username'), hashed_pw, request.form.get('name'), request.form.get('role')))
            conn.commit()
            log_system_audit("User Created", current_user.name, f"Created new user '{request.form.get('username')}' with role '{request.form.get('role')}'.")
        except: pass 
        return redirect(url_for('officers'))
    cursor.execute("SELECT * FROM users ORDER BY id ASC")
    users_list = cursor.fetchall()
    conn.close()
    return render_template('officers.html', users=users_list)

@app.route('/delete_officer/<int:id>', methods=['POST'])
@login_required
def delete_officer(id):
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT username FROM users WHERE id = %s", (id,))
    deleted_user = cursor.fetchone()
    cursor.execute("DELETE FROM users WHERE id = %s AND id != %s", (id, current_user.id))
    conn.commit()
    if deleted_user:
        log_system_audit("User Deleted", current_user.name, f"Deleted user '{deleted_user['username']}'.")
    conn.close()
    return redirect(url_for('officers'))

@app.route('/reset_password/<int:id>', methods=['POST'])
@login_required
def reset_password(id):
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    
    new_password = request.form.get('new_password')
    if not new_password:
        flash("Password cannot be empty.", "error")
        return redirect(url_for('officers'))

    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT username FROM users WHERE id = %s", (id,))
    target_user = cursor.fetchone()
    
    if target_user:
        hashed_pw = generate_password_hash(new_password)
        cursor.execute("UPDATE users SET password = %s WHERE id = %s", (hashed_pw, id))
        conn.commit()
        log_system_audit("Password Reset", current_user.name, f"Admin reset password for user '{target_user['username']}'.")
        flash(f"Password for {target_user['username']} reset successfully!", "success")
    else:
        flash("User not found.", "error")

    conn.close()
    return redirect(url_for('officers'))

@app.route('/force_unlock/<int:id>', methods=['POST'])
@login_required
def force_unlock(id):
    if current_user.role != 'Admin': return redirect(url_for('inspection_queue'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE tablets SET status = 'Pending', inspected_by = NULL WHERE id = %s", (id,))
    conn.commit()
    log_device_history(id, "Force Unlocked", current_user.name, "Pending", "Admin force unlocked the device")
    log_system_audit("Device Force Unlocked", current_user.name, f"Admin unlocked tablet T-{id}.")
    conn.close()
    return redirect(url_for('inspection_queue'))

@app.route('/undo_inspection/<int:id>', methods=['POST'])
@login_required
def undo_inspection(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT status, inspected_by FROM tablets WHERE id = %s", (id,))
    tablet = cursor.fetchone()
    
    if not tablet:
        conn.close()
        return redirect(url_for('dashboard'))
        
    if current_user.role != 'Admin' and current_user.name != tablet['inspected_by']:
        flash("🚫 Access Denied! You can only undo your own inspections.", "error")
        conn.close()
        return redirect(request.referrer or url_for('dashboard'))
        
    cursor.execute("UPDATE tablets SET status = 'Pending', inspected_by = NULL, inspection_data = NULL, battery_drain_time = '-' WHERE id = %s", (id,))
    conn.commit()
    
    log_device_history(id, "Inspection Undone", current_user.name, "Pending", "User reversed the inspection decision. Moved back to Queue.")
    conn.close()
    
    flash(f"✅ Device T-{id} inspection undone! Back in the Inspection Queue.", "success")
    return redirect(url_for('inspection_queue'))

@app.route('/search')
@login_required
def search():
    query = request.args.get('q', '').strip()
    if not query: return jsonify([])
    conn = get_db_connection()
    cursor = conn.cursor()
    search_term = f"%{query}%"
    cursor.execute("SELECT id, brand, model, serial_number, asset_no, imei_number, status, inspected_by FROM tablets WHERE serial_number LIKE %s OR imei_number LIKE %s OR asset_no LIKE %s OR CONCAT('T-', id) LIKE %s LIMIT 5", (search_term, search_term, search_term, search_term))
    results = cursor.fetchall()
    conn.close()
    return jsonify(results)

@app.route('/reports', methods=['GET'])
@login_required
def reports():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    district = request.args.get('district', '')
    brand = request.args.get('brand', '')
    status = request.args.get('status', '')
    inspector = request.args.get('inspector', '')
    
    page = request.args.get('page', 1, type=int)
    per_page = 50
    offset = (page - 1) * per_page
    
    cursor.execute("SELECT name FROM users ORDER BY name ASC")
    users = cursor.fetchall()
    
    base_query = " FROM tablets WHERE status NOT IN ('Pending', 'Locked') AND is_deleted = 0"
    params = []
    
    if district and district != 'All Districts':
        base_query += " AND district = %s"
        params.append(district)
    if brand and brand != 'All Brands':
        base_query += " AND brand = %s"
        params.append(brand)
    if status and status != 'All Statuses':
        base_query += " AND status = %s"
        params.append(status)
    if inspector and inspector != 'All Inspectors':
        base_query += " AND inspected_by = %s"
        params.append(inspector)
        
    count_query = "SELECT COUNT(*) as total" + base_query
    cursor.execute(count_query, params)
    total_records = cursor.fetchone()['total']
    
    total_pages = (total_records + per_page - 1) // per_page
    if total_pages == 0: 
        total_pages = 1
        
    query = "SELECT *" + base_query + " ORDER BY id DESC LIMIT %s OFFSET %s"
    paginated_params = params + [per_page, offset]
    
    cursor.execute(query, paginated_params)
    tablets = cursor.fetchall()
    
    for t in tablets:
        t['details'] = {}
        if t['inspection_data']:
            try: t['details'] = json.loads(t['inspection_data'])
            except: pass
    conn.close()
    
    return render_template('reports.html', tablets=tablets, req=request.args, users=users, page=page, total_pages=total_pages, total_records=total_records)

@app.route('/history')
@login_required
def history():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT 
            DATE_FORMAT(registered_at, '%Y-%m') as batch_month,
            COUNT(*) as total,
            SUM(CASE WHEN status IN ('Passed', 'Minor Issues', 'Defective', 'Dead Device (DOA)') THEN 1 ELSE 0 END) as inspected,
            SUM(CASE WHEN status = 'Minor Issues' THEN 1 ELSE 0 END) as minor_issues,
            SUM(CASE WHEN status IN ('Defective', 'Dead Device (DOA)') THEN 1 ELSE 0 END) as defective,
            SUM(CASE WHEN status = 'Passed' THEN 1 ELSE 0 END) as dispatched
        FROM tablets
        GROUP BY batch_month
        ORDER BY batch_month DESC
    """)
    batches = cursor.fetchall()
    conn.close()
    return render_template('history.html', batches=batches)

@app.route('/tablet/<int:id>')
@login_required
def tablet_timeline(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM tablets WHERE id = %s", (id,))
    tablet = cursor.fetchone()
    
    if not tablet:
        conn.close()
        return redirect(url_for('dashboard'))
        
    if tablet['inspection_data']:
        try:
            tablet['details'] = json.loads(tablet['inspection_data'])
        except:
            tablet['details'] = {}
    else:
        tablet['details'] = {}
        
    cursor.execute("SELECT * FROM device_history WHERE tablet_id = %s ORDER BY timestamp DESC", (id,))
    history_logs = cursor.fetchall()
    conn.close()
    
    return render_template('timeline.html', tablet=tablet, history=history_logs)

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    conn = get_db_connection()
    cursor = conn.cursor()

    if request.method == 'POST':
        for key, val in request.form.items():
            if key.startswith('target_'):
                dist_name = key.replace('target_', '')
                try:
                    t_val = int(val)
                    cursor.execute("UPDATE district_targets SET target_count = %s WHERE district_name = %s", (t_val, dist_name))
                except:
                    pass
        
        new_target = request.form.get('batch_target')
        if new_target:
            cursor.execute("UPDATE settings SET batch_target = %s WHERE id = 1", (new_target,))
            
        conn.commit()
        log_system_audit("Settings Updated", current_user.name, "Admin updated System Settings & District Targets.")
        flash("Settings and District Targets updated successfully!", "success")
        return redirect(url_for('settings'))

    cursor.execute("SELECT batch_target FROM settings WHERE id = 1")
    res = cursor.fetchone()
    target = res['batch_target'] if res else 540
    
    cursor.execute("SELECT * FROM district_targets ORDER BY district_name ASC")
    district_targets = cursor.fetchall()
    
    conn.close()
    return render_template('settings.html', batch_target=target, district_targets=district_targets)

@app.route('/audit')
@login_required
def audit(): 
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM system_audit ORDER BY timestamp DESC LIMIT 100")
    logs = cursor.fetchall()
    conn.close()
    return render_template('audit.html', logs=logs)

@app.route('/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_tablet(id):
    if current_user.role != 'Admin': 
        return redirect(url_for('dashboard'))
        
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        d = request.form
        raw_sn = d.get('serial_number', '').strip()
        raw_imei = d.get('imei_number', '').strip()
        asset_no = d.get('asset_no', '').strip()
        
        imei = clean_imei(raw_imei)
        if not is_valid_imei(imei):
            flash("⛔ Update Failed: IMEI Number must be EXACTLY 15 digits!", "error")
            return redirect(url_for('edit_tablet', id=id))

        sn = format_and_validate_sn(raw_sn)
        if not sn:
            flash("⛔ Update Failed: Serial Number must be at least 8 alphanumeric characters.", "error")
            return redirect(url_for('edit_tablet', id=id))

        if asset_no and not is_valid_asset_no(asset_no):
            flash("⛔ Update Failed: Asset Number must be EXACTLY 5 digits! Only numbers allowed.", "error")
            return redirect(url_for('edit_tablet', id=id))

        cursor.execute("SELECT id FROM tablets WHERE (serial_number = %s OR imei_number = %s) AND id != %s", (sn, imei, id))
        if cursor.fetchone():
            flash(f"⛔ Serial Number '{sn}' or IMEI '{imei}' already exists in another record!", "error")
            return redirect(url_for('edit_tablet', id=id))
            
        district = d.get('district')
        brand = d.get('brand')
        model = d.get('model', '')
        charger = d.get('charger_status')
        cable = d.get('cable_status')
        simpin = d.get('simpin_status')
        doc = d.get('doc_status')
        battery_drain_time = d.get('battery_drain_time', '-')
       
        verdict = d.get('verdict')
        inspection_data_json = None
        if verdict:
            checklist = {k: v for k, v in d.items() if k not in ['district', 'brand', 'model', 'serial_number', 'asset_no', 'imei_number', 'charger_status', 'cable_status', 'simpin_status', 'doc_status', 'verdict', 'battery_drain_time']}
            inspection_data_json = json.dumps(checklist)
            
            sql = """UPDATE tablets SET 
                     district=%s, brand=%s, model=%s, serial_number=%s, asset_no=%s, imei_number=%s, 
                     charger_status=%s, cable_status=%s, simpin_status=%s, doc_status=%s,
                     status=%s, inspection_data=%s, battery_drain_time=%s 
                     WHERE id=%s"""
            cursor.execute(sql, (district, brand, model, sn, asset_no, imei, charger, cable, simpin, doc, verdict, inspection_data_json, battery_drain_time, id))
        else:
            sql = """UPDATE tablets SET 
                     district=%s, brand=%s, model=%s, serial_number=%s, asset_no=%s, imei_number=%s, 
                     charger_status=%s, cable_status=%s, simpin_status=%s, doc_status=%s, battery_drain_time=%s 
                     WHERE id=%s"""
            cursor.execute(sql, (district, brand, model, sn, asset_no, imei, charger, cable, simpin, doc, battery_drain_time, id))
            
        conn.commit()
        
        log_device_history(id, "Data Edited", current_user.name, verdict, "Admin updated device details and inspection data.")
        log_system_audit("Device Edited", current_user.name, f"Admin updated data for tablet T-{id}.")
        conn.close()
        
        flash("✅ Device data updated successfully!", "success")
        return redirect(url_for('tablet_timeline', id=id))
        
    cursor.execute("SELECT * FROM tablets WHERE id = %s", (id,))
    tablet = cursor.fetchone()
    conn.close()
    if not tablet: return redirect(url_for('dashboard'))
    
    if tablet['inspection_data']:
        try:
            tablet['details'] = json.loads(tablet['inspection_data'])
        except:
            tablet['details'] = {}
    else:
        tablet['details'] = {}
        
    return render_template('edit.html', tablet=tablet)

@app.route('/delete/<int:id>', methods=['POST'])
@login_required
def delete_tablet(id):
    if current_user.role != 'Admin': 
        return redirect(url_for('dashboard'))
        
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE tablets SET is_deleted = 1 WHERE id = %s", (id,))
        conn.commit()
        log_system_audit("Device Deleted", current_user.name, f"Admin deleted tablet T-{id} from the system permanently.")
        flash(f"🗑️ Device T-{id} deleted successfully.", "success")
    except Exception as e:
        flash(f"❌ Error deleting device: {str(e)}", "error")
    finally:
        conn.close()
        
    return redirect(request.referrer or url_for('reports'))

@app.route('/nuke_ghosts')
@login_required
def nuke_ghosts():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SET FOREIGN_KEY_CHECKS=0;")
        cursor.execute("DELETE FROM device_history WHERE tablet_id IN (1, 2, 3);")
        cursor.execute("DELETE FROM tablets WHERE id IN (1, 2, 3);")
        cursor.execute("SET FOREIGN_KEY_CHECKS=1;")
        conn.commit()
        flash("💥 Ghost records (T-1, T-2, T-3) nuked successfully!", "success")
    except Exception as e:
        flash(f"❌ Error nuking ghosts: {str(e)}", "error")
    finally:
        conn.close()
    return redirect(url_for('dashboard'))

@app.route('/performance')
@login_required
def performance():
    if current_user.role != 'Admin': 
        return redirect(url_for('dashboard'))
        
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT name, role FROM users")
    users = cursor.fetchall()

    cursor.execute("SELECT registered_by, COUNT(id) as count FROM tablets GROUP BY registered_by")
    reg_data = {row['registered_by']: row['count'] for row in cursor.fetchall() if row['registered_by']}

    cursor.execute("""
        SELECT inspected_by,
               COUNT(id) as total_inspected,
               SUM(CASE WHEN status = 'Passed' THEN 1 ELSE 0 END) as passed,
               SUM(CASE WHEN status = 'Minor Issues' THEN 1 ELSE 0 END) as minor,
               SUM(CASE WHEN status IN ('Defective', 'Dead Device (DOA)') THEN 1 ELSE 0 END) as defective
        FROM tablets
        WHERE inspected_by IS NOT NULL AND inspected_by != ''
        GROUP BY inspected_by
    """)
    insp_data = {row['inspected_by']: row for row in cursor.fetchall()}
    conn.close()

    perf_data = []
    
    max_reg = 0
    max_defects = 0

    for u in users:
        name = u['name']
        r_count = reg_data.get(name, 0)
        i_info = insp_data.get(name, {'total_inspected':0, 'passed':0, 'minor':0, 'defective':0})
        
        defects = int(i_info['defective'] or 0)
        
        if r_count > max_reg: max_reg = r_count
        if defects > max_defects: max_defects = defects

        total_activity = r_count + int(i_info['total_inspected'] or 0)
        
        perf_data.append({
            'name': name,
            'role': u['role'],
            'registered': r_count,
            'inspected': int(i_info['total_inspected'] or 0),
            'passed': int(i_info['passed'] or 0),
            'minor': int(i_info['minor'] or 0),
            'defective': defects,
            'total_activity': total_activity,
            'badges': [] 
        })

    for user in perf_data:
        if user['registered'] == max_reg and max_reg > 0:
            user['badges'].append('Speedster ⚡')
        if user['defective'] == max_defects and max_defects > 0:
            user['badges'].append('Eagle Eye 🦅')

    perf_data.sort(key=lambda x: x['total_activity'], reverse=True)

    return render_template('performance.html', perf_data=perf_data)

@app.route('/export_performance_summary')
@login_required
def export_performance_summary():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT name, role FROM users")
    users = cursor.fetchall()
    cursor.execute("SELECT registered_by, COUNT(id) as count FROM tablets GROUP BY registered_by")
    reg_data = {row['registered_by']: row['count'] for row in cursor.fetchall() if row['registered_by']}
    cursor.execute("""
        SELECT inspected_by, COUNT(id) as total_inspected,
               SUM(CASE WHEN status = 'Passed' THEN 1 ELSE 0 END) as passed,
               SUM(CASE WHEN status = 'Minor Issues' THEN 1 ELSE 0 END) as minor,
               SUM(CASE WHEN status IN ('Defective', 'Dead Device (DOA)') THEN 1 ELSE 0 END) as defective
        FROM tablets WHERE inspected_by IS NOT NULL AND inspected_by != '' GROUP BY inspected_by
    """)
    insp_data = {row['inspected_by']: row for row in cursor.fetchall()}
    conn.close()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Staff Name', 'Role', 'Registered Count', 'Inspected Count', 'Passed', 'Minor Issues', 'Defective', 'Total Activity'])
    
    for u in users:
        n = u['name']
        r_count = reg_data.get(n, 0)
        i = insp_data.get(n, {'total_inspected':0, 'passed':0, 'minor':0, 'defective':0})
        total = r_count + int(i['total_inspected'] or 0)
        writer.writerow([n, u['role'], r_count, i['total_inspected'], i['passed'], i['minor'], i['defective'], total])
    
    return Response(output.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment;filename=Tabcore_Staff_Performance.csv"})

@app.route('/generate_handover_pdf', methods=['GET'])
@login_required
def generate_handover_pdf():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                        Paragraph, Spacer, KeepTogether)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from datetime import datetime
        import pytz, io, re

        district = request.args.get('district', '').strip()
        brand    = request.args.get('brand', '')
        status   = request.args.get('status', '')

        conn = get_db_connection(); cursor = conn.cursor()
        query = """SELECT id, serial_number, asset_no, imei_number, brand, model,
                          district, status, charger_status, cable_status,
                          simpin_status
                   FROM tablets
                   WHERE status NOT IN ('Pending', 'Locked')
                   AND is_deleted = 0"""
        params = []
        if district and district != 'All Districts':
            query += " AND district = %s"; params.append(district)
        if brand and brand != 'All Brands':
            query += " AND brand = %s"; params.append(brand)
        if status and status != 'All Statuses':
            query += " AND status = %s"; params.append(status)
        query += " ORDER BY district ASC, id ASC"
        cursor.execute(query, params)
        tablets = cursor.fetchall()
        conn.close()

        if not tablets:
            return "No inspected tablets found for the selected filter."

        # ── Time & Doc ID ─────────────────────────────────────────────
        tz           = pytz.timezone('Asia/Colombo')
        now          = datetime.now(tz)
        doc_id       = f"HO-{now.strftime('%Y%m%d-%H%M%S')}"
        display_time = now.strftime('%Y-%m-%d  %I:%M %p')

        # ── Counts ────────────────────────────────────────────────────
        total     = len(tablets)
        passed    = sum(1 for t in tablets if t['status'] == 'Passed')
        minor     = sum(1 for t in tablets if t['status'] == 'Minor Issues')
        defective = sum(1 for t in tablets if t['status'] == 'Defective')
        doa_count = sum(1 for t in tablets if t['status'] == 'Dead Device (DOA)')

        miss_chg = sum(1 for t in tablets if str(t.get('charger_status','')).lower() == 'missing')
        dmg_chg  = sum(1 for t in tablets if str(t.get('charger_status','')).lower() == 'damaged')
        miss_cbl = sum(1 for t in tablets if str(t.get('cable_status','')).lower()   == 'missing')
        dmg_cbl  = sum(1 for t in tablets if str(t.get('cable_status','')).lower()   == 'damaged')
        miss_pin = sum(1 for t in tablets if str(t.get('simpin_status','')).lower()  == 'missing')

        # ── Dynamic District Label ─────────────────
        show_district_column = True
        if district and district != 'All Districts':
            dist_label = district.upper()
            show_district_column = False
        else:
            unique_districts = list(set([str(t.get('district', '')).strip().upper() for t in tablets if str(t.get('district', '')).strip()]))
            if len(unique_districts) == 1:
                dist_label = unique_districts[0]
                show_district_column = False
            elif len(unique_districts) > 1:
                dist_label = "MULTIPLE DISTRICTS"
            else:
                dist_label = "ALL DISTRICTS"

        # 🔴 Dynamic Brand & Model Extractor (Space + Hyphen Fix)
        bm_set = set()
        for t in tablets:
            b = str(t.get('brand') or '').strip().capitalize()
            # ඉර දෙපැත්තේ තියෙන හිස්තැන් ඔක්කොම මකනවා (උදා: TB - 8505X -> TB-8505X)
            m = str(t.get('model') or '').strip().upper()
            m = re.sub(r'\s*-\s*', '-', m) 
            
            if b or m:
                combined = re.sub(' +', ' ', f"{b} {m}").strip()
                bm_set.add(combined)
        
        models_str = " | ".join(sorted(list(bm_set))) if bm_set else "N/A"

        # ── PDF Setup ─────────────────────────────────────────────────
        buffer     = io.BytesIO()
        # 🔴 ප්‍රින්ට් කරද්දී යට කැපෙන නිසා bottomMargin එක 55 ඉඳන් 65 ට වැඩි කළා
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=28, leftMargin=28,
                                topMargin=95,  bottomMargin=65)
        elements = []
        styles   = getSampleStyleSheet()

        small = ParagraphStyle('small', parent=styles['Normal'],
                               fontSize=8, leading=11, fontName='Helvetica')
        bold9 = ParagraphStyle('bold9', parent=styles['Normal'],
                               fontSize=9, fontName='Helvetica-Bold')
        red7  = ParagraphStyle('red7',  parent=styles['Normal'],
                               fontSize=7.5, fontName='Helvetica-Oblique',
                               textColor=colors.HexColor("#c0392b"))

        # ── Header / Footer ───────────────────────────────────────────
        def draw_hf(c, d):
            c.saveState()
            W = d.pagesize[0]; H = d.pagesize[1]

            c.setFillColor(colors.HexColor("#1a2a6c"))
            c.rect(0, H-78, W, 78, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#f0a500"))
            c.rect(0, H-82, W, 4, fill=1, stroke=0)

            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawCentredString(W/2, H-32, "DEPARTMENT OF CENSUS AND STATISTICS")
            c.setFont("Helvetica-Bold", 9)
            c.drawCentredString(W/2, H-50, f"ICT DIVISION  |  STORES HANDOVER REPORT  —  {dist_label}")
            c.setFont("Helvetica", 8)
            c.setFillColor(colors.HexColor("#aaccee"))
            c.drawCentredString(W/2, H-66, f"Document Ref: {doc_id}   |   Date & Time: {display_time}")

            # 🔴 අළු පාට කොටුව තව ටිකක් උස කරලා අකුරු ටික උඩට ගත්තා ප්‍රින්ටර් එකේ කැපෙන්නේ නැති වෙන්න (y = 14 ඉඳන් 25 ට ගෙනාවා)
            c.setFillColor(colors.HexColor("#eeeeee"))
            c.rect(0, 0, W, 45, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#555555"))
            c.setFont("Helvetica-Oblique", 7.5)
            c.drawString(28, 22, "TABCORE System  |  Dept. of Census & Statistics, Sri Lanka" f"  |  {display_time}")
            c.drawRightString(W-28, 22, f"Page {d.page}")
            c.restoreState()

        # ── 1. Meta row ───────────────────────────────────────────────
        meta_data = [
            [
                Paragraph(f"<b>Document Ref:</b>  {doc_id}", small),
                Paragraph(f"<b>Handover Date & Time:</b>  {display_time}", small),
                Paragraph(f"<b>Prepared By:</b>  {current_user.name}", small)
            ],
            [
                Paragraph(f"<b>Included Device Models:</b>  <font color='#1a2a6c'><b>{models_str}</b></font>", small),
                "", ""
            ]
        ]
        meta = Table(meta_data, colWidths=[180, 190, 168])
        meta.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#dddddd")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#f7f9ff")),
            ('TOPPADDING',   (0,0),(-1,-1), 6),
            ('BOTTOMPADDING',(0,0),(-1,-1), 6),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
            ('SPAN',         (0,1),(-1,1)),
        ]))
        elements.append(meta)
        elements.append(Spacer(1, 10))

        # ── 2. Summary cards ──────────────────────────────────────────
        def sumcell(label, val, col="#1a2a6c"):
            return Paragraph(f"<font size='14' color='{col}'><b>{val}</b></font><br/><font size='7' color='#555555'>{label}</font>", small)

        s_cards = Table([[
            sumcell("Total Devices", total,    "#1a2a6c"),
            sumcell("Passed ✅",     passed,    "#006600"),
            sumcell("Minor ⚠️",      minor,     "#cc7700"),
            sumcell("Defective ❌",  defective, "#cc0000"),
            sumcell("DOA ☠",        doa_count, "#6b21a8"),
        ]], colWidths=[107]*5)
        s_cards.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 1,   colors.HexColor("#1a2a6c")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#f0f4ff")),
            ('ALIGN',        (0,0),(-1,-1), 'CENTER'),
            ('TOPPADDING',   (0,0),(-1,-1), 8),
            ('BOTTOMPADDING',(0,0),(-1,-1), 8),
        ]))
        elements.append(Paragraph("<b>INSPECTION & ACCESSORY EXCEPTION SUMMARY</b>", bold9))
        elements.append(Spacer(1, 4))
        elements.append(s_cards)
        elements.append(Spacer(1, 5))

        acc_row = Table([[
            Paragraph(f"<b>Chargers</b> — Missing: <font color='red'>{miss_chg}</font>  |  Damaged: <font color='#cc7700'>{dmg_chg}</font>", small),
            Paragraph(f"<b>Data Cables</b> — Missing: <font color='red'>{miss_cbl}</font>  |  Damaged: <font color='#cc7700'>{dmg_cbl}</font>", small),
            Paragraph(f"<b>SIM Pins</b> — Missing: <font color='red'>{miss_pin}</font>", small),
        ]], colWidths=[178, 178, 179])
        acc_row.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#dddddd")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#fffaf5")),
            ('TOPPADDING',   (0,0),(-1,-1), 6),
            ('BOTTOMPADDING',(0,0),(-1,-1), 6),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
        ]))
        elements.append(acc_row)
        elements.append(Spacer(1, 15))

        # ── 3. Main device table ──────────────────────────────────────
        elements.append(Paragraph("<b>DEVICE HANDOVER REGISTER</b>", bold9))
        elements.append(Spacer(1, 4))

        if show_district_column:
            hdr = [['No.', 'T-ID', 'Asset No', 'District', 'Serial Number', 'IMEI Number', 'Status', 'Accessories']]
            col_widths = [25, 35, 50, 60, 100, 100, 60, 108]
            status_col_idx = 6
            acc_col_idx = 7
        else:
            hdr = [['No.', 'T-ID', 'Asset No', 'Serial Number', 'IMEI Number', 'Status', 'Accessories']]
            col_widths = [30, 45, 60, 110, 110, 65, 118]
            status_col_idx = 5
            acc_col_idx = 6

        STATUS_BG = {
            'Passed':            colors.HexColor("#d4edda"),
            'Minor Issues':      colors.HexColor("#fff3cd"),
            'Defective':         colors.HexColor("#f8d7da"),
            'Dead Device (DOA)': colors.HexColor("#ede9fe"),
        }
        STATUS_FG = {
            'Passed':            colors.HexColor("#155724"),
            'Minor Issues':      colors.HexColor("#856404"),
            'Defective':         colors.HexColor("#721c24"),
            'Dead Device (DOA)': colors.HexColor("#4c1d95"),
        }
        STATUS_LBL = {
            'Passed':            'Passed',
            'Minor Issues':      'Minor Issues',
            'Defective':         'Defective',
            'Dead Device (DOA)': 'DOA ☠',
        }

        style_cmds = [
            ('BACKGROUND',    (0,0),(-1,0),  colors.HexColor("#1a2a6c")),
            ('TEXTCOLOR',     (0,0),(-1,0),  colors.white),
            ('FONTNAME',      (0,0),(-1,0),  'Helvetica-Bold'),
            ('FONTSIZE',      (0,0),(-1,-1), 7),
            ('ALIGN',         (0,0),(-1,-1), 'CENTER'),
            ('GRID',          (0,0),(-1,-1), 0.5, colors.grey),
            ('VALIGN',        (0,0),(-1,-1), 'MIDDLE'),
            ('TOPPADDING',    (0,0),(-1,-1), 5),
            ('BOTTOMPADDING', (0,0),(-1,-1), 5),
            ('ROWBACKGROUNDS',(0,1),(-1,-1), [colors.white, colors.HexColor("#fafafa")]),
        ]

        rows = []
        for ri, t in enumerate(tablets, start=1):
            mis, dam = [], []
            for key, lbl in [('charger_status','Chg'), ('cable_status', 'Cbl'), ('simpin_status', 'Pin')]:
                v = str(t.get(key,'')).lower().strip()
                if   'missing' in v: mis.append(lbl)
                elif 'damage'  in v: dam.append(lbl)
            if not mis and not dam:
                acc = "Full Set"
            else:
                parts = []
                if mis: parts.append("No: " + ",".join(mis))
                if dam: parts.append("Dmg: " + ",".join(dam))
                acc = " | ".join(parts)

            st  = t['status']
            lbl = STATUS_LBL.get(st, st)

            if st in STATUS_BG:
                style_cmds.append(('BACKGROUND', (status_col_idx, ri), (status_col_idx, ri), STATUS_BG[st]))
                style_cmds.append(('TEXTCOLOR',  (status_col_idx, ri), (status_col_idx, ri), STATUS_FG[st]))
                style_cmds.append(('FONTNAME',   (status_col_idx, ri), (status_col_idx, ri), 'Helvetica-Bold'))

            if acc != "Full Set":
                style_cmds.append(('BACKGROUND', (acc_col_idx, ri), (acc_col_idx, ri), colors.HexColor("#fff8e1")))
                style_cmds.append(('TEXTCOLOR',  (acc_col_idx, ri), (acc_col_idx, ri), colors.HexColor("#7a4900")))

            row_data = [
                str(ri),
                f"T-{t['id']}",
                t.get('asset_no') or '-'
            ]
            
            if show_district_column:
                row_data.append(t.get('district') or '-')
                
            row_data.extend([
                t['serial_number'],
                t.get('imei_number') or '-',
                lbl,
                acc
            ])
            
            rows.append(row_data)

        main_tbl = Table(hdr + rows, colWidths=col_widths, repeatRows=1)
        main_tbl.setStyle(TableStyle(style_cmds))
        elements.append(main_tbl)
        elements.append(Spacer(1, 6))
        elements.append(Paragraph("* Accessories Legend: Chg = Charger  |  Cbl = Data Cable  |  Pin = SIM Pin", red7))

        # ── 4. Remarks + Signatures ───────────────────────────────────
        sig_block = KeepTogether([
            Spacer(1, 14),
            Paragraph("<b>Remarks / Notes:</b>", small),
            Spacer(1, 5),
            Paragraph("................................................................................................................................................................", small),
            Spacer(1, 4),
            Paragraph("................................................................................................................................................................", small),
            Spacer(1, 28),

            Table([[
                Table([
                    [Paragraph("<b>Handed Over By (ICT):</b>", small)],
                    [Spacer(1, 14)],
                    [Paragraph("Name:  .........................................", small)],
                    [Spacer(1, 12)],
                    [Paragraph("Sign & Date:  ..................................", small)],
                ], colWidths=[238], style=[('TOPPADDING',(0,0),(-1,-1),0), ('BOTTOMPADDING',(0,0),(-1,-1),2)]),

                Paragraph("", small),

                Table([
                    [Paragraph("<b>Received By (Stores):</b>", small)],
                    [Spacer(1, 14)],
                    [Paragraph("Name:  .........................................", small)],
                    [Spacer(1, 12)],
                    [Paragraph("Sign & Date:  ..................................", small)],
                    [Spacer(1, 12)],
                    [Paragraph("Rubber Stamp:", small)],
                    [Table([[""]], colWidths=[110], rowHeights=[45], style=[('BOX',(0,0),(-1,-1), 0.8, colors.grey)])],
                ], colWidths=[238], style=[('TOPPADDING',(0,0),(-1,-1),0), ('BOTTOMPADDING',(0,0),(-1,-1),2)]),

            ]], colWidths=[240, 15, 240], style=[('VALIGN',(0,0),(-1,-1),'TOP'), ('TOPPADDING',(0,0),(-1,-1),0), ('BOTTOMPADDING',(0,0),(-1,-1),0)]),
        ])
        elements.append(sig_block)

        doc.build(elements, onFirstPage=draw_hf, onLaterPages=draw_hf)
        buffer.seek(0)
        return Response(buffer, mimetype='application/pdf', headers={"Content-Disposition": f"attachment;filename=Stores_Handover_{doc_id}.pdf"})

    except Exception as e:
        import traceback
        return (f"Error: {str(e)}<br><pre>{traceback.format_exc()}</pre>")

from flask import Response

@app.route('/generate_full_inspection_pdf', methods=['GET'])
@login_required
def generate_full_inspection_pdf():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                        Paragraph, Spacer, PageBreak, KeepTogether)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from datetime import datetime
        import pytz, io, re, json

        district = request.args.get('district', '').strip()
        brand    = request.args.get('brand', '')
        status   = request.args.get('status', '')

        conn = get_db_connection()
        cursor = conn.cursor()
        query = """SELECT * FROM tablets
                   WHERE status NOT IN ('Pending', 'Locked')
                   AND is_deleted = 0"""
        params = []
        if district and district != 'All Districts':
            query += " AND district = %s"; params.append(district)
        if brand and brand != 'All Brands':
            query += " AND brand = %s"; params.append(brand)
        if status and status != 'All Statuses':
            query += " AND status = %s"; params.append(status)
        query += " ORDER BY district ASC, id ASC"
        cursor.execute(query, params)
        tablets = cursor.fetchall()
        conn.close()

        if not tablets:
            return "No inspected tablets found for this filter!"

        # ── Time & Doc ID ─────────────────────────────────
        tz           = pytz.timezone('Asia/Colombo')
        now          = datetime.now(tz)
        doc_id       = f"FI-{now.strftime('%Y%m%d-%H%M%S')}"
        display_time = now.strftime('%Y-%m-%d  %I:%M %p')

        # ── Counts ─────────────────────────────────────────
        total     = len(tablets)
        passed    = sum(1 for t in tablets if t['status'] == 'Passed')
        minor     = sum(1 for t in tablets if t['status'] == 'Minor Issues')
        defective = sum(1 for t in tablets if t['status'] == 'Defective')
        doa_count = sum(1 for t in tablets if t['status'] == 'Dead Device (DOA)')

        miss_chg = sum(1 for t in tablets if str(t.get('charger_status','')).lower() == 'missing')
        dmg_chg  = sum(1 for t in tablets if str(t.get('charger_status','')).lower() == 'damaged')
        miss_cbl = sum(1 for t in tablets if str(t.get('cable_status','')).lower()   == 'missing')
        dmg_cbl  = sum(1 for t in tablets if str(t.get('cable_status','')).lower()   == 'damaged')
        miss_pin = sum(1 for t in tablets if str(t.get('simpin_status','')).lower()  == 'missing')

        # ── District label ─────────────────────────────────
        show_district_col = True
        unique_dists = list(set([
            str(t.get('district','')).strip().title()
            for t in tablets
            if str(t.get('district','')).strip()
        ]))
        if district and district != 'All Districts':
            dist_label = district.title()
            show_district_col = False
        elif len(unique_dists) == 1:
            dist_label = unique_dists[0].upper()
            show_district_col = False
        elif len(unique_dists) > 1:
            dist_label = "MULTIPLE DISTRICTS"
        else:
            dist_label = "ALL DISTRICTS"

        # ── Device Models ──────────────────────────────────
        bm_set = set()
        for t in tablets:
            b = str(t.get('brand') or '').strip().title()
            m = str(t.get('model') or '').strip().upper()
            m = re.sub(r'\s*-\s*', '-', m)
            combined = re.sub(' +', ' ', f"{b} {m}").strip()
            if combined: bm_set.add(combined)
        models_str = " | ".join(sorted(bm_set)) if bm_set else "N/A"

        # ── District breakdown ─────────────────────────────
        dist_map = {}
        for t in tablets:
            d = str(t.get('district') or 'Unknown').strip().title()
            if d not in dist_map:
                dist_map[d] = {'total':0,'passed':0,'minor':0,'defective':0,'doa':0}
            dist_map[d]['total'] += 1
            if   t['status'] == 'Passed':             dist_map[d]['passed']    += 1
            elif t['status'] == 'Minor Issues':        dist_map[d]['minor']     += 1
            elif t['status'] == 'Defective':           dist_map[d]['defective'] += 1
            elif t['status'] == 'Dead Device (DOA)':   dist_map[d]['doa']       += 1

        # ── PDF Setup ──────────────────────────────────────
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=landscape(A4),
            rightMargin=20, leftMargin=20,
            topMargin=92, bottomMargin=52
        )
        elements = []
        styles   = getSampleStyleSheet()

        small = ParagraphStyle('small', parent=styles['Normal'],
                               fontSize=8, leading=11, fontName='Helvetica')
        bold9 = ParagraphStyle('bold9', parent=styles['Normal'],
                               fontSize=9,  fontName='Helvetica-Bold')
        red7  = ParagraphStyle('red7',  parent=styles['Normal'],
                               fontSize=7.5, fontName='Helvetica-Oblique',
                               textColor=colors.HexColor("#c0392b"))

        # ── Part tracker (shared via list so closure can mutate) ──────
        part_tracker = ["PART 1: IDENTITY & SUMMARY"]

        def draw_hf(c, d):
            c.saveState()
            W = d.pagesize[0]; H = d.pagesize[1]

            # Navy header bar
            c.setFillColor(colors.HexColor("#1a2a6c"))
            c.rect(0, H-78, W, 78, fill=1, stroke=0)
            # Gold accent line
            c.setFillColor(colors.HexColor("#f0a500"))
            c.rect(0, H-82, W, 4, fill=1, stroke=0)

            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawCentredString(W/2, H-30,
                "DEPARTMENT OF CENSUS AND STATISTICS")
            c.setFont("Helvetica-Bold", 9)
            # Use the part label stored in canvas object (set per-page)
            part_lbl = getattr(c, '_tabcore_part', part_tracker[0])
            c.drawCentredString(W/2, H-50,
                f"FULL INSPECTION REPORT  —  {dist_label}  |  {part_lbl}")
            c.setFont("Helvetica", 7.5)
            c.setFillColor(colors.HexColor("#aaccee"))
            c.drawCentredString(W/2, H-65,
                f"Document Ref: {doc_id}   |   Date & Time: {display_time}"
                f"   |   Generated By: {current_user.name}")

            # Footer bar
            c.setFillColor(colors.HexColor("#eeeeee"))
            c.rect(0, 0, W, 38, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#555555"))
            c.setFont("Helvetica-Oblique", 7.5)
            c.drawString(20, 15,
                "TABCORE System  |  Dept. of Census & Statistics, Sri Lanka  |"
                "  Full Inspection Export")
            c.drawRightString(W-20, 15, f"Page {d.page}")
            c.restoreState()

        # ── Helper: set part label per page ───────────────
        # We use a "beforeDrawPage" trick via a custom flowable
        from reportlab.platypus import Flowable

        class SetPartLabel(Flowable):
            def __init__(self, label):
                super().__init__()
                self.label = label
                self.width = self.height = 0
            def draw(self):
                self.canv._tabcore_part = self.label
                part_tracker[0] = self.label

        # ── Summary section (first page only) ─────────────
        meta = Table([[
            Paragraph(f"<b>Document Ref:</b>  {doc_id}", small),
            Paragraph(f"<b>Report Date & Time:</b>  {display_time}", small),
            Paragraph(f"<b>Generated By:</b>  {current_user.name}", small),
        ]], colWidths=[266, 267, 267])
        meta.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#dddddd")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#f7f9ff")),
            ('TOPPADDING',   (0,0),(-1,-1), 6),
            ('BOTTOMPADDING',(0,0),(-1,-1), 6),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
        ]))

        models_para = Paragraph(
            f"<b>Device Models:</b>  <font color='#1a2a6c'><b>{models_str}</b></font>",
            small)
        meta2 = Table([[models_para]], colWidths=[800])
        meta2.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#f7f9ff")),
            ('TOPPADDING',   (0,0),(-1,-1), 5),
            ('BOTTOMPADDING',(0,0),(-1,-1), 5),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
        ]))

        def sumcell(label, val, col="#1a2a6c"):
            return Paragraph(
                f"<font size='16' color='{col}'><b>{val}</b></font><br/>"
                f"<font size='7' color='#555555'>{label}</font>", small)

        s_cards = Table([[
            sumcell("Total Devices", total,    "#1a2a6c"),
            sumcell("Passed ✅",     passed,    "#006600"),
            sumcell("Minor ⚠",       minor,     "#cc7700"),
            sumcell("Defective ❌",  defective, "#cc0000"),
            sumcell("DOA",           doa_count, "#6b21a8"),
        ]], colWidths=[160]*5)
        s_cards.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 1,   colors.HexColor("#1a2a6c")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#f0f4ff")),
            ('ALIGN',        (0,0),(-1,-1), 'CENTER'),
            ('TOPPADDING',   (0,0),(-1,-1), 9),
            ('BOTTOMPADDING',(0,0),(-1,-1), 9),
        ]))

        acc_row = Table([[
            Paragraph(f"<b>Chargers</b> — Missing: <font color='red'>{miss_chg}</font>"
                      f"  |  Damaged: <font color='#cc7700'>{dmg_chg}</font>", small),
            Paragraph(f"<b>Data Cables</b> — Missing: <font color='red'>{miss_cbl}</font>"
                      f"  |  Damaged: <font color='#cc7700'>{dmg_cbl}</font>", small),
            Paragraph(f"<b>SIM Pins</b> — Missing: <font color='red'>{miss_pin}</font>",
                      small),
        ]], colWidths=[266, 267, 267])
        acc_row.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#dddddd")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#fffaf5")),
            ('TOPPADDING',   (0,0),(-1,-1), 6),
            ('BOTTOMPADDING',(0,0),(-1,-1), 6),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
        ]))

        # Set Part 1 label before summary
        elements.append(SetPartLabel("PART 1: IDENTITY & SUMMARY"))
        elements.append(meta)
        elements.append(Spacer(1, 4))
        elements.append(meta2)
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(
            "<b>GLOBAL INSPECTION & ACCESSORY SUMMARY</b>", bold9))
        elements.append(Spacer(1, 4))
        elements.append(s_cards)
        elements.append(Spacer(1, 5))
        elements.append(acc_row)

        # District breakdown (if multiple)
        if len(dist_map) > 1:
            elements.append(Spacer(1, 8))
            elements.append(Paragraph("<b>DISTRICT BREAKDOWN</b>", bold9))
            elements.append(Spacer(1, 4))
            d_hdr  = [['District','Total','Passed','Minor','Defective','DOA']]
            d_rows = [[dn, dc['total'], dc['passed'], dc['minor'],
                       dc['defective'], dc['doa']]
                      for dn, dc in sorted(dist_map.items())]
            d_tbl  = Table(d_hdr + d_rows,
                           colWidths=[140, 60, 65, 60, 70, 60])
            d_tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0,0),(-1,0),  colors.HexColor("#1a2a6c")),
                ('TEXTCOLOR',     (0,0),(-1,0),  colors.white),
                ('FONTNAME',      (0,0),(-1,0),  'Helvetica-Bold'),
                ('FONTSIZE',      (0,0),(-1,-1), 8),
                ('ALIGN',         (0,0),(-1,-1), 'CENTER'),
                ('ALIGN',         (0,1),(0,-1),  'LEFT'),
                ('LEFTPADDING',   (0,1),(0,-1),  8),
                ('GRID',          (0,0),(-1,-1), 0.5, colors.grey),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),
                 [colors.white, colors.HexColor("#f5f5f5")]),
                ('TOPPADDING',    (0,0),(-1,-1), 4),
                ('BOTTOMPADDING', (0,0),(-1,-1), 4),
            ]))
            elements.append(d_tbl)

        elements.append(Spacer(1, 14))

        # ── CHUNKING — smarter: fill each page pair properly ──────────
        # Part 1: ~14 rows first page, ~16 subsequent
        # Part 2: ~16 rows per page
        # We keep Part1 chunk == Part2 chunk per "batch"
        ROWS_PER_PAGE = 15

        batches = []
        idx = 0
        first = True
        while idx < len(tablets):
            cap = (13 if first else ROWS_PER_PAGE)
            batches.append(tablets[idx: idx + cap])
            idx  += cap
            first = False

        current_serial = 1

        for batch_idx, batch in enumerate(batches):

            # ── PART 1 table ─────────────────────────────────────────
            elements.append(Paragraph(
                f"<b>PART 1: DEVICE IDENTITY & SUMMARY</b>"
                f"  <font size='8' color='#555555'>"
                f"(Devices {current_serial} – {current_serial+len(batch)-1}"
                f" of {total})</font>", bold9))
            elements.append(Spacer(1, 5))

            if show_district_col:
                p1_hdr  = ['No.','T-ID','Asset No','District',
                           'Serial Number','IMEI Number',
                           'Inspector','Insp. Date','Battery Drain','Status']
                p1_cols = [30, 42, 58, 72, 112, 112, 118, 78, 82, 72]
            else:
                p1_hdr  = ['No.','T-ID','Asset No',
                           'Serial Number','IMEI Number',
                           'Inspector','Insp. Date','Battery Drain','Status']
                p1_cols = [30, 42, 62, 128, 130, 140, 90, 90, 78]

            p1_data = [p1_hdr]
            p1_cmds = [
                ('BACKGROUND',    (0,0),(-1,0),  colors.HexColor("#1a2a6c")),
                ('TEXTCOLOR',     (0,0),(-1,0),  colors.white),
                ('FONTNAME',      (0,0),(-1,0),  'Helvetica-Bold'),
                ('FONTSIZE',      (0,0),(-1,0),  8),
                ('FONTSIZE',      (0,1),(-1,-1), 7.5),
                ('ALIGN',         (0,0),(-1,-1), 'CENTER'),
                ('VALIGN',        (0,0),(-1,-1), 'MIDDLE'),
                ('GRID',          (0,0),(-1,-1), 0.5, colors.grey),
                ('TOPPADDING',    (0,0),(-1,-1), 5),
                ('BOTTOMPADDING', (0,0),(-1,-1), 5),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),
                 [colors.white, colors.HexColor("#f7f9ff")]),
            ]

            STATUS_FG = {
                'Passed':            colors.HexColor("#15803d"),
                'Minor Issues':      colors.HexColor("#b45309"),
                'Defective':         colors.HexColor("#b91c1c"),
                'Dead Device (DOA)': colors.HexColor("#4c1d95"),
            }
            STATUS_BG_ROW = {
                'Dead Device (DOA)': colors.HexColor("#f5f0ff"),
                'Defective':         colors.HexColor("#fff8f8"),
            }

            for ri, t in enumerate(batch, 1):
                sn = current_serial + ri - 1

                inspector = t.get('inspected_by') or '-'
                insp_date = '-'
                rd = t.get('registered_at')
                if rd:
                    try: insp_date = str(rd)[:10]
                    except: pass

                drain_raw = str(t.get('battery_drain_time') or '').strip()
                if drain_raw and drain_raw not in ('-','','None','0'):
                    try:
                        mins = int(''.join(filter(str.isdigit, drain_raw)))
                        if mins > 0:
                            rating = (f"{mins} min ✓" if mins > 60 else
                                      f"{mins} min ~" if mins >= 45 else
                                      f"{mins} min ✗")
                        else:
                            rating = 'Not Tested'
                    except:
                        rating = drain_raw
                else:
                    rating = 'Not Tested'

                st = t['status']

                # Status color
                if st in STATUS_FG:
                    p1_cmds.append(('TEXTCOLOR', (-1,ri),(-1,ri), STATUS_FG[st]))
                    p1_cmds.append(('FONTNAME',  (-1,ri),(-1,ri), 'Helvetica-Bold'))

                # Row tint for DOA/Defective
                if st in STATUS_BG_ROW:
                    p1_cmds.append(('BACKGROUND',(0,ri),(-2,ri), STATUS_BG_ROW[st]))

                row = [str(sn), f"T-{t['id']}",
                       str(t.get('asset_no') or '-')]
                if show_district_col:
                    row.append(str(t.get('district') or '-').title())
                row.extend([
                    t['serial_number'],
                    str(t.get('imei_number') or '-'),
                    inspector, insp_date, rating, st
                ])
                p1_data.append(row)

            p1_tbl = Table(p1_data, colWidths=p1_cols, repeatRows=1)
            p1_tbl.setStyle(TableStyle(p1_cmds))
            elements.append(p1_tbl)

            elements.append(KeepTogether([
                Spacer(1, 6),
                Paragraph(
                    "<font size=7.5 color='#555555'>"
                    "* Battery Drain: ✓ = Good (&gt;60 min) &nbsp;"
                    "~ = Moderate (45–60 min) &nbsp; ✗ = Poor (&lt;45 min)"
                    "</font>", styles['Normal'])
            ]))

            # ── PAGE BREAK + PART 2 ───────────────────────────────────
            elements.append(PageBreak())
            elements.append(SetPartLabel("PART 2: HARDWARE & ACCESSORY DETAILS"))

            elements.append(Paragraph(
                f"<b>PART 2: HARDWARE & ACCESSORY DETAILS</b>"
                f"  <font size='8' color='#555555'>"
                f"(Devices {current_serial} – {current_serial+len(batch)-1}"
                f" of {total})</font>", bold9))
            elements.append(Spacer(1, 5))

            p2_hdr  = ['No.','T-ID','S/N',
                       'Dsp','Tch','Bat','Cam','WiF','BT',
                       'GPS','Spk','Mic','Pwr','SIM','Prt',
                       'Accessories','Notes']
            p2_cols = [28, 40, 88,
                       29,29,29,29,29,29,
                       29,29,29,29,29,29,
                       110, 168]

            p2_data = [p2_hdr]
            p2_cmds = [
                ('BACKGROUND',    (0,0),(-1,0),  colors.HexColor("#1a2a6c")),
                ('TEXTCOLOR',     (0,0),(-1,0),  colors.white),
                ('FONTNAME',      (0,0),(-1,0),  'Helvetica-Bold'),
                ('FONTSIZE',      (0,0),(-1,0),  7.5),
                ('FONTSIZE',      (0,1),(-1,-1), 6.5),
                ('ALIGN',         (0,0),(-1,-1), 'CENTER'),
                ('VALIGN',        (0,0),(-1,-1), 'MIDDLE'),
                ('GRID',          (0,0),(-1,-1), 0.5, colors.grey),
                ('TOPPADDING',    (0,0),(-1,-1), 4),
                ('BOTTOMPADDING', (0,0),(-1,-1), 4),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),
                 [colors.white, colors.HexColor("#f7f9ff")]),
            ]

            HW_KEYS = ['display','touch','battery','cameras','wifi','bt',
                       'gps','speaker','mic','p_btn','sim','charging']
            HW_BG = {
                'P':   colors.HexColor("#d4edda"),
                'M':   colors.HexColor("#fff3cd"),
                'F':   colors.HexColor("#f8d7da"),
                'N/A': colors.HexColor("#e2e8f0"),
            }

            for ri, t in enumerate(batch, 1):
                sn = current_serial + ri - 1
                det = {}
                if t.get('inspection_data'):
                    try: det = json.loads(t['inspection_data'])
                    except: pass

                hw_vals = []
                for ci, k in enumerate(HW_KEYS, 3):
                    v    = str(det.get(k, '-')).lower()
                    char = ('P'   if v == 'pass'               else
                            'M'   if v in ['minor','partial']   else
                            'F'   if v == 'fail'                else
                            'N/A' if v == 'n/a'                 else '-')
                    hw_vals.append(char)
                    if char in HW_BG:
                        p2_cmds.append(('BACKGROUND',
                                        (ci,ri),(ci,ri), HW_BG[char]))

                mis, dam = [], []
                for key, lbl in [('charger_status','Chg'),
                                  ('cable_status',  'Cbl'),
                                  ('simpin_status', 'Pin')]:
                    v = str(t.get(key,'')).lower().strip()
                    if   'missing' in v: mis.append(lbl)
                    elif 'damage'  in v: dam.append(lbl)

                if not mis and not dam:
                    acc = "Full Set"
                else:
                    parts = []
                    if mis: parts.append("No: " + ",".join(mis))
                    if dam: parts.append("Dmg: " + ",".join(dam))
                    acc = " | ".join(parts)
                    p2_cmds.append(('BACKGROUND',
                                    (15,ri),(15,ri),
                                    colors.HexColor("#fff5f5")))
                    p2_cmds.append(('TEXTCOLOR',
                                    (15,ri),(15,ri),
                                    colors.HexColor("#7a0000")))

                notes = str(det.get('inspector_notes','') or '-')[:55]

                p2_data.append(
                    [str(sn), f"T-{t['id']}", t['serial_number'][:13]]
                    + hw_vals + [acc, notes]
                )

            p2_tbl = Table(p2_data, colWidths=p2_cols, repeatRows=1)
            p2_tbl.setStyle(TableStyle(p2_cmds))
            elements.append(p2_tbl)

            legend = (
                "<font size=7.5 color='#555555'>"
                "<b>Hardware:</b> Dsp=Display | Tch=Touch | Bat=Battery | "
                "Cam=Camera | WiF=Wi-Fi | BT=Bluetooth | GPS | "
                "Spk=Speaker | Mic | Pwr=Power Btn | SIM | Prt=Charging Port"
                "&nbsp;&nbsp;<b>Verdict:</b> P=Pass &nbsp; M=Minor &nbsp; "
                "F=Fail &nbsp; N/A=Not Applicable"
                "&nbsp;&nbsp;<b>Acc:</b> Chg=Charger, Cbl=Cable, Pin=SIM Pin"
                "</font>"
            )
            elements.append(KeepTogether([
                Spacer(1, 6),
                Paragraph(legend, styles['Normal'])
            ]))

            current_serial += len(batch)

            # Page break between batches (not after last)
            if batch_idx < len(batches) - 1:
                elements.append(PageBreak())
                elements.append(SetPartLabel("PART 1: IDENTITY & SUMMARY"))

        # ── Build ──────────────────────────────────────────
        doc.build(elements,
                  onFirstPage=draw_hf,
                  onLaterPages=draw_hf)
        buffer.seek(0)
        return Response(
            buffer, mimetype='application/pdf',
            headers={"Content-Disposition":
                     f"attachment;filename={dist_label}_Full_Inspection_{doc_id}.pdf"})

    except Exception as e:
        import traceback
        return (f"Error: {str(e)}<br>"
                f"<pre>{traceback.format_exc()}</pre>")
                
@app.route('/generate_defect_report_pdf', methods=['GET'])
@login_required
def generate_defect_report_pdf():
    if current_user.role != 'Admin': return redirect(url_for('dashboard'))
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                        Paragraph, Spacer, KeepTogether,
                                        PageBreak)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from datetime import datetime
        import pytz, io, re, json

        district = request.args.get('district', '').strip()
        brand    = request.args.get('brand', '')

        conn = get_db_connection()
        cursor = conn.cursor()
        query = """SELECT * FROM tablets
                   WHERE status IN ('Minor Issues','Defective','Dead Device (DOA)')
                   AND is_deleted = 0"""
        params = []
        if district and district != 'All Districts':
            query += " AND district = %s"; params.append(district)
        if brand and brand != 'All Brands':
            query += " AND brand = %s"; params.append(brand)
        query += " ORDER BY status ASC, district ASC, id ASC"
        cursor.execute(query, params)
        tablets = cursor.fetchall()
        conn.close()

        if not tablets:
            return ("No defective or flagged devices found for this filter!"
                    " All devices passed inspection.")

        # ── Time & IDs ────────────────────────────────────────────────
        tz           = pytz.timezone('Asia/Colombo')
        now          = datetime.now(tz)
        doc_id       = f"DR-{now.strftime('%Y%m%d-%H%M%S')}"
        display_time = now.strftime('%Y-%m-%d  %I:%M %p')

        # ── Counts ────────────────────────────────────────────────────
        total     = len(tablets)
        minor     = sum(1 for t in tablets if t['status'] == 'Minor Issues')
        defective = sum(1 for t in tablets if t['status'] == 'Defective')
        doa_count = sum(1 for t in tablets if t['status'] == 'Dead Device (DOA)')

        miss_chg = sum(1 for t in tablets if str(t.get('charger_status','')).lower() == 'missing')
        dmg_chg  = sum(1 for t in tablets if str(t.get('charger_status','')).lower() == 'damaged')
        miss_cbl = sum(1 for t in tablets if str(t.get('cable_status','')).lower()   == 'missing')
        dmg_cbl  = sum(1 for t in tablets if str(t.get('cable_status','')).lower()   == 'damaged')
        miss_pin = sum(1 for t in tablets if str(t.get('simpin_status','')).lower()  == 'missing')

        # ── District label & column ───────────────────────────────────
        show_dist_col = True
        unique_dists  = list(set([
            str(t.get('district','')).strip().title()
            for t in tablets if str(t.get('district','')).strip()
        ]))
        if district and district != 'All Districts':
            dist_label    = district.title()
            show_dist_col = False
        elif len(unique_dists) == 1:
            dist_label    = unique_dists[0].upper()
            show_dist_col = False
        else:
            dist_label = "MULTIPLE DISTRICTS" if unique_dists else "ALL DISTRICTS"

        # ── Models ────────────────────────────────────────────────────
        bm_set = set()
        for t in tablets:
            b = str(t.get('brand') or '').strip().title()
            m = re.sub(r'\s*-\s*', '-', str(t.get('model') or '').strip().upper())
            c = re.sub(' +', ' ', f"{b} {m}").strip()
            if c: bm_set.add(c)
        models_str = " | ".join(sorted(bm_set)) or "N/A"

        # ── District breakdown ────────────────────────────────────────
        dist_map = {}
        for t in tablets:
            d = str(t.get('district') or 'Unknown').strip().title()
            if d not in dist_map:
                dist_map[d] = {'minor':0,'defective':0,'doa':0,'total':0}
            dist_map[d]['total'] += 1
            if   t['status'] == 'Minor Issues':        dist_map[d]['minor']     += 1
            elif t['status'] == 'Defective':           dist_map[d]['defective'] += 1
            elif t['status'] == 'Dead Device (DOA)':   dist_map[d]['doa']       += 1

        # ── PDF Setup ─────────────────────────────────────────────────
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=landscape(A4),
            rightMargin=20, leftMargin=20,
            topMargin=92,  bottomMargin=52
        )
        elements = []
        styles   = getSampleStyleSheet()

        small = ParagraphStyle('small', parent=styles['Normal'],
                               fontSize=8, leading=11, fontName='Helvetica')
        bold9 = ParagraphStyle('bold9', parent=styles['Normal'],
                               fontSize=9, fontName='Helvetica-Bold')
        cell_s = ParagraphStyle('cell_s', parent=styles['Normal'],
                                fontSize=7.5, leading=10, fontName='Helvetica')

        # ── Header / Footer ───────────────────────────────────────────
        def draw_hf(c, d):
            c.saveState()
            W = d.pagesize[0]; H = d.pagesize[1]

            # Deep red header
            c.setFillColor(colors.HexColor("#7f1d1d"))
            c.rect(0, H-78, W, 78, fill=1, stroke=0)
            # Gold accent
            c.setFillColor(colors.HexColor("#f0a500"))
            c.rect(0, H-82, W, 4, fill=1, stroke=0)

            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawCentredString(W/2, H-30,
                "DEPARTMENT OF CENSUS AND STATISTICS")
            c.setFont("Helvetica-Bold", 9)
            c.drawCentredString(W/2, H-50,
                f"DEFECT & SHORTAGE ACTION REPORT  —  {dist_label}")
            c.setFont("Helvetica", 7.5)
            c.setFillColor(colors.HexColor("#fecaca"))
            c.drawCentredString(W/2, H-65,
                f"Document Ref: {doc_id}   |   Date & Time: {display_time}"
                f"   |   Generated By: {current_user.name}")

            # Footer
            c.setFillColor(colors.HexColor("#eeeeee"))
            c.rect(0, 0, W, 38, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#555555"))
            c.setFont("Helvetica-Oblique", 7.5)
            c.drawString(20, 15,
                "TABCORE System  |  Dept. of Census & Statistics, Sri Lanka"
                "  |  Defect & Action Report Export")
            c.drawRightString(W-20, 15, f"Page {d.page}")
            c.restoreState()

        # ── Meta row ──────────────────────────────────────────────────
        meta = Table([[
            Paragraph(f"<b>Document Ref:</b>  {doc_id}", small),
            Paragraph(f"<b>Report Date & Time:</b>  {display_time}", small),
            Paragraph(f"<b>Generated By:</b>  {current_user.name}", small),
        ]], colWidths=[266, 267, 267])
        meta.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#dddddd")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#fff5f5")),
            ('TOPPADDING',   (0,0),(-1,-1), 6),
            ('BOTTOMPADDING',(0,0),(-1,-1), 6),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
        ]))

        meta2 = Table([[
            Paragraph(
                f"<b>Included Models:</b>  "
                f"<font color='#7f1d1d'><b>{models_str}</b></font>", small)
        ]], colWidths=[800])
        meta2.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#fff5f5")),
            ('TOPPADDING',   (0,0),(-1,-1), 5),
            ('BOTTOMPADDING',(0,0),(-1,-1), 5),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
        ]))

        # ── Summary cards ─────────────────────────────────────────────
        def sumcell(label, val, col="#7f1d1d"):
            return Paragraph(
                f"<font size='16' color='{col}'><b>{val}</b></font><br/>"
                f"<font size='7' color='#555555'>{label}</font>", small)

        s_cards = Table([[
            sumcell("Total Flagged", total,    "#7f1d1d"),
            sumcell("Minor Issues ⚠", minor,   "#cc7700"),
            sumcell("Defective ❌",  defective, "#cc0000"),
            sumcell("DOA ☠",        doa_count, "#6b21a8"),
        ]], colWidths=[200]*4)
        s_cards.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 1,   colors.HexColor("#7f1d1d")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#dddddd")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#fff0f0")),
            ('ALIGN',        (0,0),(-1,-1), 'CENTER'),
            ('TOPPADDING',   (0,0),(-1,-1), 9),
            ('BOTTOMPADDING',(0,0),(-1,-1), 9),
        ]))

        acc_row = Table([[
            Paragraph(
                f"<b>Chargers</b> — Missing: <font color='red'>{miss_chg}</font>"
                f"  |  Damaged: <font color='#cc7700'>{dmg_chg}</font>", small),
            Paragraph(
                f"<b>Data Cables</b> — Missing: <font color='red'>{miss_cbl}</font>"
                f"  |  Damaged: <font color='#cc7700'>{dmg_cbl}</font>", small),
            Paragraph(
                f"<b>SIM Pins</b> — Missing: <font color='red'>{miss_pin}</font>",
                small),
        ]], colWidths=[266, 267, 267])
        acc_row.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1), 0.5, colors.HexColor("#aaaaaa")),
            ('INNERGRID',    (0,0),(-1,-1), 0.5, colors.HexColor("#dddddd")),
            ('BACKGROUND',   (0,0),(-1,-1), colors.HexColor("#fffaf5")),
            ('TOPPADDING',   (0,0),(-1,-1), 6),
            ('BOTTOMPADDING',(0,0),(-1,-1), 6),
            ('LEFTPADDING',  (0,0),(-1,-1), 8),
        ]))

        elements.append(meta)
        elements.append(Spacer(1, 4))
        elements.append(meta2)
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(
            "<b>DEFECT & ACCESSORY SHORTAGE SUMMARY</b>", bold9))
        elements.append(Spacer(1, 4))
        elements.append(s_cards)
        elements.append(Spacer(1, 5))
        elements.append(acc_row)

        # ── District breakdown (if multiple) ──────────────────────────
        if len(dist_map) > 1:
            elements.append(Spacer(1, 8))
            elements.append(Paragraph(
                "<b>DISTRICT DEFECT BREAKDOWN</b>", bold9))
            elements.append(Spacer(1, 4))
            d_hdr  = [['District','Total Flagged','Minor','Defective','DOA']]
            d_rows = [[dn, dc['total'], dc['minor'],
                       dc['defective'], dc['doa']]
                      for dn, dc in sorted(dist_map.items())]
            d_tbl  = Table(d_hdr + d_rows,
                           colWidths=[160, 90, 80, 90, 80])
            d_tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0,0),(-1,0),  colors.HexColor("#7f1d1d")),
                ('TEXTCOLOR',     (0,0),(-1,0),  colors.white),
                ('FONTNAME',      (0,0),(-1,0),  'Helvetica-Bold'),
                ('FONTSIZE',      (0,0),(-1,-1), 8),
                ('ALIGN',         (0,0),(-1,-1), 'CENTER'),
                ('ALIGN',         (0,1),(0,-1),  'LEFT'),
                ('LEFTPADDING',   (0,1),(0,-1),  8),
                ('GRID',          (0,0),(-1,-1), 0.5, colors.grey),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),
                 [colors.white, colors.HexColor("#fff5f5")]),
                ('TOPPADDING',    (0,0),(-1,-1), 4),
                ('BOTTOMPADDING', (0,0),(-1,-1), 4),
            ]))
            elements.append(d_tbl)

        elements.append(Spacer(1, 14))

        # ── Defect details helper ─────────────────────────────────────
        HW_MAP = {
            'display':  'Display',
            'touch':    'Touch',
            'battery':  'Battery',
            'cameras':  'Camera',
            'wifi':     'Wi-Fi',
            'bt':       'Bluetooth',
            'gps':      'GPS',
            'speaker':  'Speaker',
            'mic':      'Mic',
            'p_btn':    'Power Btn',
            'sim':      'SIM Slot',
            'charging': 'Charging Port',
        }
        
        # 🔴 වෙනස්කම: DOA වලටත් Send to Repair හැදුවා
        ACTION_MAP = {
            'Minor Issues':      'Monitor / Note',
            'Defective':         'Send to Repair',
            'Dead Device (DOA)': 'Send to Repair',
        }

        def build_defect_cell(t):
            det = {}
            if t.get('inspection_data'):
                try: det = json.loads(t['inspection_data'])
                except: pass

            lines = []

            # Hardware fails
            hw_fails = []
            for k, lbl in HW_MAP.items():
                v = str(det.get(k, '')).lower()
                if v == 'fail':
                    hw_fails.append(
                        f"<font color='#b91c1c'><b>{lbl}</b>: FAIL</font>")
                elif v in ['minor', 'partial']:
                    hw_fails.append(
                        f"<font color='#b45309'><b>{lbl}</b>: Minor</font>")
            if hw_fails:
                lines.append("HW: " + " | ".join(hw_fails))

            # Accessory issues
            acc_issues = []
            for key, lbl in [('charger_status', 'Charger'),
                              ('cable_status',   'Cable'),
                              ('simpin_status',  'SIM Pin')]:
                v = str(t.get(key, '')).lower().strip()
                if   'missing' in v:
                    acc_issues.append(
                        f"<font color='#b91c1c'><b>{lbl}</b>: Missing</font>")
                elif 'damage' in v:
                    acc_issues.append(
                        f"<font color='#b45309'><b>{lbl}</b>: Damaged</font>")
            if acc_issues:
                lines.append("Acc: " + " | ".join(acc_issues))

            # Battery drain
            drain = str(t.get('battery_drain_time') or '').strip()
            if drain and drain not in ('-', '', 'None', '0'):
                try:
                    mins = int(''.join(filter(str.isdigit, drain)))
                    if mins > 0 and mins < 45:
                        lines.append(
                            f"<font color='#b91c1c'><b>Battery Drain:</b>"
                            f" {mins} min (Poor)</font>")
                    elif mins > 0:
                        lines.append(
                            f"<b>Battery Drain:</b> {mins} min")
                except: pass

            # DOA special
            if t['status'] == 'Dead Device (DOA)':
                lines = [
                    "<font color='#4c1d95'><b>DEAD ON ARRIVAL</b> — "
                    "No power response. Completely unresponsive.</font>"
                ]

            # Inspector notes (full, up to 120 chars)
            notes = str(det.get('inspector_notes', '') or '').strip()
            if notes and notes != '-':
                n = notes[:120] + ('...' if len(notes) > 120 else '')
                lines.append(
                    f"<i><font color='#555555'>Note: {n}</font></i>")

            if not lines:
                lines = ["<font color='#888888'>No specific defect recorded."
                         " Check inspector notes.</font>"]

            return Paragraph("<br/>".join(lines), cell_s)

        # ── Main defect table ─────────────────────────────────────────
        elements.append(Paragraph("<b>DEFECTIVE DEVICES REGISTER</b>", bold9))
        elements.append(Spacer(1, 5))

        if show_dist_col:
            t_hdr  = ['No.','T-ID','Asset','District','Serial Number',
                      'IMEI','Inspector','Insp. Date','Status',
                      'Action Required','Defects & Shortages']
            t_cols = [28, 40, 52, 65, 95, 92, 65, 62, 62, 72, 167]
            st_idx  = 8
            act_idx = 9
            def_idx = 10
        else:
            t_hdr  = ['No.','T-ID','Asset','Serial Number',
                      'IMEI','Inspector','Insp. Date','Status',
                      'Action Required','Defects & Shortages']
            t_cols = [28, 40, 58, 110, 105, 80, 72, 72, 82, 153]
            st_idx  = 7
            act_idx = 8
            def_idx = 9

        t_data = [t_hdr]
        t_cmds = [
            ('BACKGROUND',    (0,0),(-1,0),  colors.HexColor("#7f1d1d")),
            ('TEXTCOLOR',     (0,0),(-1,0),  colors.white),
            ('FONTNAME',      (0,0),(-1,0),  'Helvetica-Bold'),
            ('FONTSIZE',      (0,0),(-1,0),  7.5),
            ('FONTSIZE',      (0,1),(-1,-1), 7),
            ('ALIGN',         (0,0),(def_idx-1,-1), 'CENTER'),
            ('ALIGN',         (def_idx,0),(def_idx,-1), 'LEFT'),
            ('VALIGN',        (0,0),(-1,-1), 'MIDDLE'),
            ('GRID',          (0,0),(-1,-1), 0.5, colors.grey),
            ('TOPPADDING',    (0,0),(-1,-1), 5),
            ('BOTTOMPADDING', (0,0),(-1,-1), 5),
            ('TOPPADDING',    (0,0),(-1,0),  6),
            ('BOTTOMPADDING', (0,0),(-1,0),  6),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),
             [colors.white, colors.HexColor("#fff8f8")]),
        ]

        STATUS_FG = {
            'Minor Issues':      colors.HexColor("#b45309"),
            'Defective':         colors.HexColor("#b91c1c"),
            'Dead Device (DOA)': colors.HexColor("#4c1d95"),
        }
        STATUS_ROW_BG = {
            'Dead Device (DOA)': colors.HexColor("#f5f0ff"),
            'Defective':         colors.HexColor("#fff5f5"),
        }
        
        # 🔴 වෙනස්කම: Supplier Return අයින් කළා
        ACTION_BG = {
            'Monitor / Note':    colors.HexColor("#fff8e0"),
            'Send to Repair':    colors.HexColor("#ffe4e4"),
        }
        ACTION_FG = {
            'Monitor / Note':    colors.HexColor("#92400e"),
            'Send to Repair':    colors.HexColor("#7f1d1d"),
        }

        for ri, t in enumerate(tablets, 1):
            st     = t['status']
            action = ACTION_MAP.get(st, 'Review')

            # Status colors
            if st in STATUS_FG:
                t_cmds.append(('TEXTCOLOR',
                               (st_idx,ri),(st_idx,ri), STATUS_FG[st]))
                t_cmds.append(('FONTNAME',
                               (st_idx,ri),(st_idx,ri), 'Helvetica-Bold'))

            # Row tint
            if st in STATUS_ROW_BG:
                t_cmds.append(('BACKGROUND',
                               (0,ri),(def_idx-1,ri), STATUS_ROW_BG[st]))

            # Action cell color
            if action in ACTION_BG:
                t_cmds.append(('BACKGROUND',
                               (act_idx,ri),(act_idx,ri), ACTION_BG[action]))
                t_cmds.append(('TEXTCOLOR',
                               (act_idx,ri),(act_idx,ri), ACTION_FG[action]))
                t_cmds.append(('FONTNAME',
                               (act_idx,ri),(act_idx,ri), 'Helvetica-Bold'))

            # Inspector & date
            inspector = str(t.get('inspected_by') or '-')
            insp_date = '-'
            rd = t.get('registered_at')
            if rd:
                try: insp_date = str(rd)[:10]
                except: pass

            row = [str(ri), f"T-{t['id']}",
                   str(t.get('asset_no') or '-')]
            if show_dist_col:
                row.append(str(t.get('district') or '-').title())
            row.extend([
                t['serial_number'],
                str(t.get('imei_number') or '-'),
                inspector,
                insp_date,
                st,
                action,
                build_defect_cell(t),
            ])
            t_data.append(row)

        # Chunk table if many rows to avoid overflow
        ROWS_PER_PAGE = 18
        if len(t_data) <= ROWS_PER_PAGE + 1:
            tbl = Table(t_data, colWidths=t_cols, repeatRows=1)
            tbl.setStyle(TableStyle(t_cmds))
            elements.append(tbl)
        else:
            chunks  = [t_data[0:1]]
            current = []
            for r in t_data[1:]:
                current.append(r)
                if len(current) == ROWS_PER_PAGE:
                    chunks.append(t_data[0:1] + current)
                    current = []
            if current:
                chunks.append(t_data[0:1] + current)

            for ci, chunk in enumerate(chunks):
                tbl = Table(chunk, colWidths=t_cols, repeatRows=1)
                tbl.setStyle(TableStyle(t_cmds))
                elements.append(tbl)
                if ci < len(chunks) - 1:
                    elements.append(PageBreak())

        # ── Legend ────────────────────────────────────────────────────
        # 🔴 වෙනස්කම: Legend එකේ තේරුමත් හැදුවා
        elements.append(KeepTogether([
            Spacer(1, 8),
            Paragraph(
                "<font size=7.5 color='#7f1d1d'>"
                "<b>Action Legend:</b> &nbsp;"
                "<b>Monitor/Note</b> = Minor issue, deploy with caution &nbsp;|&nbsp; "
                "<b>Send to Repair</b> = ICT hardware repair required (For Defective & DOA)"
                "</font>",
                styles['Normal'])
        ]))

        # ── Build ──────────────────────────────────────────────────────
        doc.build(elements, onFirstPage=draw_hf, onLaterPages=draw_hf)
        buffer.seek(0)
        return Response(
            buffer, mimetype='application/pdf',
            headers={"Content-Disposition":
                     f"attachment;filename={dist_label}_Defect_Action_Report_{doc_id}.pdf"})

    except Exception as e:
        import traceback
        return (f"Error: {str(e)}<br>"
                f"<pre>{traceback.format_exc()}</pre>")

# TRASH BIN LOGIC (SOFT DELETE MANAGEMENT)
# ==========================================

@app.route('/trash')
@login_required
def trash():
    if current_user.role != 'Admin':
        flash('Permission denied. Admin access required.', 'error')
        return redirect(url_for('dashboard'))

    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT * FROM tablets WHERE is_deleted = 1 ORDER BY id DESC")
    deleted_tablets = cursor.fetchall()
    conn.close()
    
    return render_template('trash.html', tablets=deleted_tablets)

@app.route('/restore/<int:id>', methods=['POST'])
@login_required
def restore_tablet(id):
    if current_user.role != 'Admin':
        return redirect(url_for('dashboard'))

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("UPDATE tablets SET is_deleted = 0 WHERE id = %s", (id,))
        conn.commit()
        
        log_system_audit("Device Restored", current_user.name, f"Admin restored tablet T-{id} from Trash Bin.")
        
        flash(f"✅ Device T-{id} restored successfully! It is back in the system.", "success")
    except Exception as e:
        flash(f"❌ Error restoring device: {str(e)}", "error")
    finally:
        if 'conn' in locals():
            conn.close()

    return redirect(url_for('trash'))

@app.route('/factory_reset_db')
@login_required
def factory_reset_db():
    if current_user.role != 'Admin':
        flash("Access Denied! Only Admins can reset the database.", "error")
        return redirect(url_for('dashboard'))

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SET FOREIGN_KEY_CHECKS = 0;")
        
        cursor.execute("TRUNCATE TABLE tablets;")
        cursor.execute("TRUNCATE TABLE device_history;")
        cursor.execute("TRUNCATE TABLE inventory_logs;")
        
        cursor.execute("UPDATE bulk_inventory SET good_qty=0, defective_qty=0, remark='';")
        
        cursor.execute("SET FOREIGN_KEY_CHECKS = 1;")
        
        conn.commit()
        conn.close()
        
        flash("🔥 System Factory Reset Successful! All data cleared and reset to zero.", "success")
        return redirect(url_for('dashboard'))
        
    except Exception as e:
        flash(f"Error resetting database: {str(e)}", "error")
        return redirect(url_for('dashboard'))

# ==========================================
# 📦 BULK INVENTORY MANAGEMENT (DISTRICT-WISE)
# ==========================================

@app.route('/inventory')
@login_required
def inventory():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM bulk_inventory ORDER BY district ASC, item_name ASC")
    inventory_data = cursor.fetchall()
    
    cursor.execute("SELECT item_name, SUM(good_qty) as total_good FROM bulk_inventory GROUP BY item_name")
    totals = cursor.fetchall()
    conn.close()
    return render_template('inventory.html', inventory_data=inventory_data, totals=totals)

@app.route('/update_inventory', methods=['POST'])
@login_required
def update_inventory():
    district = request.form['district']
    item_name = request.form['item_name']
    good_qty_change = int(request.form['good_qty'])
    defective_qty_change = int(request.form['defective_qty'])
    remark = request.form.get('remark', '')

    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT good_qty, defective_qty FROM bulk_inventory WHERE district=%s AND item_name=%s", (district, item_name))
    current_stock = cursor.fetchone()

    if current_stock:
        new_good = current_stock['good_qty'] + good_qty_change
        new_defective = current_stock['defective_qty'] + defective_qty_change

        if new_good < 0:
            conn.close()
            flash(f"❌ Error: Cannot reduce GOOD stock below zero! Current good stock is {current_stock['good_qty']}.", "danger")
            return redirect(url_for('inventory'))
            
        if new_defective < 0:
            conn.close()
            flash(f"❌ Error: Cannot reduce DEFECTIVE stock below zero! Current defective stock is {current_stock['defective_qty']}.", "danger")
            return redirect(url_for('inventory'))

    cursor.execute("""
        UPDATE bulk_inventory 
        SET good_qty = good_qty + %s, 
            defective_qty = defective_qty + %s, 
            remark = %s 
        WHERE district=%s AND item_name=%s
    """, (good_qty_change, defective_qty_change, remark, district, item_name))

    cursor.execute("""
        INSERT INTO inventory_logs (username, district, item_name, good_qty_changed, defective_qty_changed, remark)
        VALUES (%s, %s, %s, %s, %s, %s)
    """, (current_user.username, district, item_name, good_qty_change, defective_qty_change, remark))

    conn.commit()
    conn.close()
    
    flash(f"✅ {district} District - Stock updated successfully! (Logged)", "success")
    return redirect(url_for('inventory'))

@app.route('/export_inventory_pdf', methods=['POST'])
@login_required
def export_inventory_pdf():
    district = request.form.get('district')
    
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM bulk_inventory WHERE district=%s ORDER BY item_name ASC", (district,))
    inv_data = cursor.fetchall()
    conn.close()

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=40, bottomMargin=30)
    elements = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('TitleStyle', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=15, spaceAfter=5)
    elements.append(Paragraph("TABCORE ENTERPRISE SYSTEM", title_style))
    elements.append(Paragraph(f"INVENTORY HANDOVER DOCUMENT - {district.upper()} DISTRICT", title_style))
    elements.append(Spacer(1, 20))

    tz = pytz.timezone('Asia/Colombo')
    current_time = datetime.now(tz).strftime('%Y-%m-%d %I:%M %p')
    elements.append(Paragraph(f"<b>Date & Time:</b> {current_time}", styles['Normal']))
    elements.append(Paragraph(f"<b>Generated By:</b> {current_user.username}", styles['Normal']))
    elements.append(Paragraph(f"<b>Handover Location:</b> {district} District Office", styles['Normal']))
    elements.append(Spacer(1, 20))

    data = [['ITEM NAME', 'GOOD (✅)', 'DEFECTIVE (❌)', 'TOTAL QTY', 'REMARKS']]
    for item in inv_data:
        good = item['good_qty']
        defective = item['defective_qty']
        data.append([item['item_name'], str(good), str(defective), str(good + defective), item['remark'] or '-'])

    table = Table(data, colWidths=[130, 70, 80, 70, 180])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e293b')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#f8fafc'), colors.white])
    ]))
    elements.append(table)

    sig_data = [['Handed Over By:', 'Received By:', 'Authorized By:'],
                ['__________________', '__________________', '__________________'],
                ['Name / Signature', 'Name / Signature', 'Name / Signature'],
                ['Date: ....................', 'Date: ....................', 'Date: ....................']]
    
    sig_table = Table(sig_data, colWidths=[170, 170, 170])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
        ('TOPPADDING', (0, 0), (-1, -1), 15)
    ]))

    elements.append(KeepTogether([Spacer(1, 60), sig_table]))
    doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"Handover_{district}_{current_time}.pdf", mimetype='application/pdf')

@app.route('/backup_db')
@login_required
def backup_db():
    if current_user.role != 'Admin': 
        return redirect(url_for('dashboard'))
        
    import subprocess
    from datetime import datetime
    import pytz
    import os
    
    # ඩේටාබේස් එකට කනෙක්ට් වෙන විස්තර .env එකෙන් ගන්නවා
    host = os.getenv('DB_HOST', 'localhost')
    user = os.getenv('DB_USER', 'root')
    password = os.getenv('DB_PASSWORD', '')
    db_name = os.getenv('DB_NAME', 'tabcore')
    
    # MySQL Dump Command එක හදනවා
    cmd = ["mysqldump", "-h", host, "-u", user, f"-p{password}", db_name]
    
    try:
        # Command එක Run කරලා SQL ෆයිල් එක ජෙනරේට් කරනවා
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()
        
        if process.returncode != 0:
            flash(f"❌ Backup failed: {stderr.decode('utf-8')}", "error")
            return redirect(url_for('settings'))
        
        # ලංකාවේ වෙලාවට නම හදනවා
        tz = pytz.timezone('Asia/Colombo')
        filename = f"Tabcore_Database_Backup_{datetime.now(tz).strftime('%Y%m%d_%H%M')}.sql"
        
        # Backup ගත්තා කියලා System Audit එකේ සේව් කරනවා
        log_system_audit("Database Backup", current_user.name, "Admin downloaded a full SQL database backup.")
        
        # SQL ෆයිල් එක ඩවුන්ලෝඩ් වෙන්න දෙනවා
        return Response(
            stdout, 
            mimetype="application/octet-stream", 
            headers={"Content-Disposition": f"attachment;filename={filename}"}
        )
    except Exception as e:
        flash(f"❌ Error generating backup: {str(e)}", "error")
        return redirect(url_for('settings'))

if __name__ == '__main__':
    create_tables() 
    create_default_admin()
    print("TabCore Server is running on Port 80...")
    from waitress import serve
    serve(app, host='0.0.0.0', port=80)